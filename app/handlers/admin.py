import asyncio
import io
import os
import zipfile
from datetime import datetime, timedelta
from aiogram import Router, F
from aiogram.filters import Command, StateFilter
from aiogram.types import Message, CallbackQuery, InlineKeyboardMarkup, InlineKeyboardButton, BufferedInputFile, ChatMemberUpdated
from aiogram.fsm.context import FSMContext
from aiogram.fsm.state import State, StatesGroup

from app.storage import (
    get_applications, update_application_status, add_seller, get_sellers,
    get_all_products, admin_delete_product, update_product, get_product_by_id,
    product_photos,
    update_seller, delete_seller, delete_user, get_orders, get_seller_reviews,
    get_order_by_id, update_order_status, update_order_fields, get_seller_orders, get_seller,
    get_orders_by_group,
    get_seller_products, add_product, decrement_stock, to_int,
    get_promos, get_promo, add_promo, delete_promo,
    get_admins, add_admin, remove_admin, is_sub_admin,
    get_couriers, add_courier, remove_courier, is_courier,
    add_audit, get_audit,
    get_cities, add_city, remove_city, get_user, get_users, set_user_field,
    shop_notify_ids, get_owner_id,
    get_blocked, mark_blocked, unmark_blocked,
)
from app.album import collect
from app.keyboards.seller import main_menu, stars_kb, menu_for, MENU_VERSION
from app.app.config.settings import settings

router = Router()

# Mahsulot qo'shishni bo'lib yuboruvchi menyu tugmalari
_MENU_BUTTONS = {
    "🛒 Market", "🛍 Bozor", "🔎 Qidirish", "🏪 Sotuvchi bo'lish", "📦 Buyurtmalarim",
    "🛍 Savatim", "❤️ Istaklarim",
    "👤 Profil", "ℹ️ Ma'lumot", "📞 Aloqa", "🛍 Do'kon (ilova)", "❌ Bekor qilish",
    # Yangi (premium) menyu yozuvlari — eski klaviaturalar bilan birga qabul qilinadi.
    "🛍 Katalog", "🔍 Qidiruv", "🧺 Savat", "❤️ Sevimlilar",
    "🛒 Sotuvchi paneli", "📊 Sotuvchi paneli", "👥 Shahrim sellerlari",
}

async def _ack(call):
    """Tugma spinnerini DARHOL o'chiradi — og'ir ishlar tugashini kutmasdan.
    Handler boshqa handlerdan qayta chaqirilganda callback allaqachon
    javoblangan bo'lishi mumkin — shunda xato bermaydi."""
    try:
        await call.answer()
    except Exception:
        pass


# ─── Robust navigatsiya: eski (48soat+) yoki rasm xabarda edit_text ishlamaydi ──
async def _admin_nav(call, text, kb):
    try:
        if call.message.photo or call.message.video or call.message.document:
            try: await call.message.delete()
            except Exception: pass
            await call.message.answer(text, parse_mode="HTML", reply_markup=kb)
        else:
            await call.message.edit_text(text, parse_mode="HTML", reply_markup=kb)
    except Exception:
        try:
            await call.message.answer(text, parse_mode="HTML", reply_markup=kb)
        except Exception:
            pass


class AdminSellerSearch(StatesGroup):
    query = State()


class AdminMsgSeller(StatesGroup):
    text = State()


class AdminBroadcast(StatesGroup):
    text = State()


class AdminPromo(StatesGroup):
    code    = State()
    percent = State()
    limit   = State()


class AdminRestartConfirm(StatesGroup):
    word = State()


def _sellers_keyboard(items):
    rows = []
    for uid, s in items:
        rows.append([InlineKeyboardButton(
            text=f"✏️ {s['shop_name']}",
            callback_data=f"aseller_{uid}"
        )])
    rows.append([InlineKeyboardButton(text="🔍 Seller qidirish", callback_data="seller_search")])
    rows.append([InlineKeyboardButton(text="🔙 Orqaga", callback_data="admin_back")])
    return InlineKeyboardMarkup(inline_keyboard=rows)


def _sellers_text(items):
    from app.storage import get_reviews, get_orders, get_all_products
    # Reyting, buyurtma va mahsulot sonini har seller uchun alohida emas, bitta
    # o'tishda yig'amiz (N+1 takroriy aylanishlar ro'yxat katta bo'lganda
    # sekinlashtirardi).
    stars_by_seller: dict = {}
    for r in get_reviews():
        stars_by_seller.setdefault(r.get("seller_id"), []).append(r.get("stars", 0))
    orders_by_seller: dict = {}
    for o in get_orders():
        sid = o.get("seller_id")
        orders_by_seller[sid] = orders_by_seller.get(sid, 0) + 1
    products_by_seller: dict = {}
    for p in get_all_products():
        sid = p.get("seller_id")
        products_by_seller[sid] = products_by_seller.get(sid, 0) + 1

    lines = ["🏪 <b>Sellerlar</b> — jami " + str(len(items)) + " ta\n"]
    for i, (uid, s) in enumerate(items, 1):
        stars = stars_by_seller.get(int(uid), [])
        rating = round(sum(stars) / len(stars), 1) if stars else 0.0
        cnt = len(stars)
        ordc = orders_by_seller.get(int(uid), 0)
        prodc = products_by_seller.get(int(uid), 0)
        card = s.get("card_number", "")
        last4 = card[-4:] if card else "—"
        lines.append(
            f"<b>{i}. {s['shop_name']}</b>\n"
            f"   👤 {s['full_name']}\n"
            f"   🏙 {s.get('city','—')}\n"
            f"   📱 {s.get('phone','—')}\n"
            f"   💳 **** {last4}   ⭐ {rating} ({cnt})\n"
            f"   📦 {prodc} mahsulot   🛒 {ordc} buyurtma"
        )
    return "\n".join(lines)


def is_owner(uid: int) -> bool:
    return uid == settings.OWNER_ID


def is_admin(uid: int) -> bool:
    # Owner yoki sub-admin
    return is_owner(uid) or is_sub_admin(uid)


def _actor_name(u) -> str:
    return u.full_name + (f" (@{u.username})" if u.username else "")


def _admin_city(uid: int):
    a = get_admins().get(str(uid))
    return a.get("city") if a else None


def _log(actor, action: str, target: str = ""):
    role = "Owner" if is_owner(actor.id) else "Admin"
    add_audit({
        "admin_id":   actor.id,
        "admin_name": _actor_name(actor),
        "role":       role,
        "action":     action,
        "target":     target,
    })


class AdminAddProduct(StatesGroup):
    seller = State()
    name        = State()
    description = State()
    price       = State()
    photo       = State()


# ─── States ──────────────────────────────────────────────────────────────────
class AdminEditShop(StatesGroup):
    waiting_value = State()

class AdminEditProduct(StatesGroup):
    waiting_field = State()
    waiting_value = State()
    waiting_photo = State()
    waiting_video = State()


# ─── Menus ───────────────────────────────────────────────────────────────────
def admin_menu_kb(uid: int):
    if is_owner(uid):
        return InlineKeyboardMarkup(inline_keyboard=[
            [InlineKeyboardButton(text="📋 Arizalar",        callback_data="admin_applications")],
            [InlineKeyboardButton(text="🏪 Sellerlar",        callback_data="admin_sellers")],
            [InlineKeyboardButton(text="📢 Ommaviy xabar / e'lon", callback_data="admin_broadcast")],
            [InlineKeyboardButton(text="📣 Guruhlarga reklama",    callback_data="admin_ads")],
            [InlineKeyboardButton(text="➕ Mahsulot qo'shish", callback_data="admin_addprod")],
            [InlineKeyboardButton(text="📦 Mahsulotlar",      callback_data="admin_products")],
            [InlineKeyboardButton(text="🎁 Promo-kodlar",     callback_data="admin_promos")],
            [InlineKeyboardButton(text="👥 Foydalanuvchilar", callback_data="admin_users")],
            [InlineKeyboardButton(text="🚫 Bloklaganlar",     callback_data="admin_blocked")],
            [InlineKeyboardButton(text="🏙 Shaharlar",        callback_data="admin_cities")],
            [InlineKeyboardButton(text="📊 Statistika",       callback_data="admin_stats")],
            [InlineKeyboardButton(text="📑 Hisobotlar",       callback_data="admin_reports")],
            [InlineKeyboardButton(text="👮 Adminlar",         callback_data="admin_admins")],
            [InlineKeyboardButton(text="🚚 Kurierlar",        callback_data="admin_couriers")],
            [InlineKeyboardButton(text="💾 Zaxira (backup)",  callback_data="admin_backup")],
            [InlineKeyboardButton(text="🔄 Restart (hammaga yangi menyu)", callback_data="admin_restart_menu")],
        ])
    # Sub-admin: faqat seller qo'shish (arizalar) + mahsulot qo'shish
    return InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="📋 Arizalar (seller qo'shish)", callback_data="admin_applications")],
        [InlineKeyboardButton(text="➕ Mahsulot qo'shish",          callback_data="admin_addprod")],
    ])


# ─── /admin ──────────────────────────────────────────────────────────────────
@router.message(Command("admin"))
async def admin_panel(message: Message):
    if not is_admin(message.from_user.id):
        await message.answer("❌ Sizda admin huquqi yo'q.")
        return
    if is_owner(message.from_user.id):
        apps = get_applications()
        pending = [a for a in apps.values() if a.get("status") == "pending"]
        text = (
            f"<b>Admin panel</b>\n"
            f"──────────────\n"
            f"Kutilayotgan arizalar: <b>{len(pending)}</b>\n"
            f"Sellerlar: <b>{len(get_sellers())}</b>\n"
            f"Sub-adminlar: <b>{len(get_admins())}</b>"
        )
    else:
        text = (
            "<b>Admin panel</b> (yordamchi)\n"
            "──────────────\n"
            "Huquqlaringiz:\n"
            "• Arizalarni tasdiqlab seller qo'shish\n"
            "• Mahsulot qo'shish"
        )
    await message.answer(text, reply_markup=admin_menu_kb(message.from_user.id), parse_mode="HTML")


@router.callback_query(F.data == "admin_back")
async def admin_back(call: CallbackQuery):
    if not is_admin(call.from_user.id): return
    await _ack(call)
    if is_owner(call.from_user.id):
        apps = get_applications()
        pending = [a for a in apps.values() if a.get("status") == "pending"]
        text = (
            f"<b>Admin panel</b>\n"
            f"──────────────\n"
            f"Kutilayotgan arizalar: <b>{len(pending)}</b>\n"
            f"Sellerlar: <b>{len(get_sellers())}</b>\n"
            f"Sub-adminlar: <b>{len(get_admins())}</b>"
        )
    else:
        text = (
            "<b>Admin panel</b> (yordamchi)\n"
            "──────────────\n"
            "• Arizalarni tasdiqlab seller qo'shish\n"
            "• Mahsulot qo'shish"
        )
    await _admin_nav(call, text, admin_menu_kb(call.from_user.id))


# ─── Arizalar ────────────────────────────────────────────────────────────────
@router.callback_query(F.data == "admin_applications")
async def show_applications(call: CallbackQuery):
    if not is_admin(call.from_user.id): return
    await _ack(call)
    pending = {uid: a for uid, a in get_applications().items() if a.get("status") == "pending"}
    if not is_owner(call.from_user.id):
        mycity = _admin_city(call.from_user.id)
        pending = {uid: a for uid, a in pending.items() if a.get("city") == mycity}
    if not pending:
        await call.message.edit_text("✅ Kutilayotgan ariza yo'q.", reply_markup=InlineKeyboardMarkup(
            inline_keyboard=[[InlineKeyboardButton(text="🔙 Orqaga", callback_data="admin_back")]]
        ))
        return
    for uid, app in pending.items():
        kb = InlineKeyboardMarkup(inline_keyboard=[[
            InlineKeyboardButton(text="✅ Tasdiqlash", callback_data=f"approve_{uid}"),
            InlineKeyboardButton(text="❌ Rad etish",  callback_data=f"reject_{uid}"),
        ]])
        card = app.get('card_number', '')
        text = (
            f"📋 <b>Ariza #{uid}</b>\n\n"
            f"👤 {app.get('full_name')}\n"
            f"🏪 Do'kon: {app.get('shop_name','—')}\n"
            f"📱 {app.get('phone')}\n"
            f"🏙 Shahar: {app.get('city','—')}\n"
            f"💳 **** {card[-4:] if card else '—'}"
        )
        await call.message.answer(text, reply_markup=kb, parse_mode="HTML")


@router.callback_query(F.data.startswith("approve_"))
async def approve_seller(call: CallbackQuery):
    if not is_admin(call.from_user.id): return
    uid = int(call.data.split("_")[1])
    app = get_applications().get(str(uid))
    if not app:
        await call.answer("Ariza topilmadi."); return
    # Spinnerni DARHOL to'xtatamiz — yozish va xabarlar keyin davom etadi.
    await call.answer("✅ Tasdiqlandi!")
    update_application_status(uid, "approved")
    add_seller(uid, {
        "user_id": uid, "full_name": app["full_name"],
        "shop_name": app.get("shop_name") or app["full_name"], "phone": app["phone"],
        "card_number": app.get("card_number", ""), "city": app.get("city", ""),
    })
    _log(call.from_user, "Seller qo'shildi",
         f"{app['full_name']} — do'kon: {app.get('shop_name','—')} (ID:{uid}, {app.get('city','—')})")
    await call.message.edit_text(f"✅ {app['full_name']} seller sifatida tasdiqlandi!")
    try:
        from app.bot.bot import bot
        await bot.send_message(
            uid,
            "<b>Tabriklaymiz — siz endi sotuvchisiz.</b>\n\n"
            "Arizangiz tasdiqlandi. Seller panelini ochish uchun /seller "
            "buyrug'ini bosing va mahsulot qo'shishni boshlang.",
            parse_mode="HTML"
        )
    except Exception: pass


@router.callback_query(F.data.startswith("reject_"))
async def reject_seller_cb(call: CallbackQuery):
    if not is_admin(call.from_user.id): return
    uid = int(call.data.split("_")[1])
    app = get_applications().get(str(uid))
    if not app:
        await call.answer("Ariza topilmadi."); return
    await call.answer("❌ Rad etildi!")
    update_application_status(uid, "rejected")
    _log(call.from_user, "Ariza rad etildi", f"{app['full_name']} (ID:{uid})")
    await call.message.edit_text(f"❌ {app['full_name']} arizasi rad etildi.")
    try:
        from app.bot.bot import bot
        await bot.send_message(uid, "❌ Seller arizangiz rad etildi.")
    except Exception: pass


# ─── To'lov chekini tasdiqlash / rad etish ───────────────────────────────────
def _speed_text(speed: str) -> str:
    """Yetkazib berish vaqti matni — admin tugma bilan tanlagan turga qarab."""
    if speed == "24h":
        return "🟡 Yetkazib berish: 24 SOAT ichida"
    return "🟢 Yetkazib berish: BUGUN (taksi pochta)"


async def _confirm_single_order(o: dict, notify_buyer: bool = True, speed: str = "today"):
    """Bitta buyurtmani 'paid' qiladi: zaxira ayiriladi, seller va kurier(lar)
    xabardor qilinadi. Xaridorga xabar notify_buyer=True bo'lsa yuboriladi —
    savat guruhida xaridorга bitta umumiy xabar borgani uchun guruhda
    notify_buyer=False bilan chaqiriladi.

    speed — admin tugma orqali tanlagan yetkazib berish vaqti ("today"|"24h").
    Mahsulot turiga qarab admin tanlaydi va xaridorга shu vaqt ko'rsatiladi."""
    oid = o["id"]
    update_order_status(oid, "paid")
    # ── Sovg'a: BITTA profilga BIR marta ──
    # Birinchi tasdiqlangan buyurtmada sovg'a beriladi va profilga yozib
    # qo'yiladi (gift_order_id) — keyingi buyurtmalarda qayta berilmaydi.
    buyer_profile = get_user(o["buyer_id"]) or {}
    gift_now = not buyer_profile.get("gift_order_id")
    if gift_now:
        set_user_field(o["buyer_id"], "gift_order_id", oid)
    # Tanlangan yetkazib berish vaqtini buyurtmaga yozamiz (keyin /orders va
    # boshqa joylarda ham ko'rsatish mumkin).
    try:
        update_order_fields(oid, {"delivery_speed": speed})
    except Exception:
        pass
    # To'lov tasdiqlandi — zaxiradan buyurtma miqdorini ayiramiz (cheksiz bo'lsa
    # tegmaydi; 0 ga yetsa mahsulot avtomatik "tugagan" bo'ladi).
    try:
        decrement_stock(o.get("product_id"), o.get("quantity", 1))
    except Exception:
        pass

    from app.bot.bot import bot
    is_pickup = o.get("delivery") == "pickup"
    dlv_map = {
        "pickup": "🚶 O'zi olib ketadi",
        "taxi": "🚕 Taksi pochta (shu bugunoq)",
        "btc": "📦 BTC Pochta", "emu": "🚀 EMU Express", "uzum": "🍊 Uzum Pochta",
    }
    dlv_label = dlv_map.get(o.get("delivery", ""), o.get("delivery", "—"))

    # Sovg'a qatorlari — faqat birinchi buyurtmada ko'rsatiladi.
    gift_line_buyer = (
        "\n🎁 Buyurtmangizga <b>kafolatlangan sovg'a</b> qo'shib yuboriladi!"
        if gift_now else ""
    )
    gift_line_seller = (
        "🎁 <b>Bu buyurtmaga sovg'a qo'shing</b> — xaridorning birinchi buyurtmasi.\n"
        if gift_now else ""
    )

    if is_pickup:
        # ── PICKUP: xaridor o'zi olib ketadi ──
        # Sellerga xaridor ma'lumotlari KO'RSATILMAYDI.
        # Xaridorga esa do'kon kontaktini beramiz — borib olib ketishi uchun.
        shop = get_seller(o["seller_id"]) or {}
        shop_name = shop.get("shop_name", "—")
        shop_phone = shop.get("phone", "—")
        shop_city = shop.get("city", "—")

        # Xaridorga — do'kon kontakti ochiladi
        if notify_buyer:
            try:
                await bot.send_message(
                    o["buyer_id"],
                    f"✅ <b>Buyurtma #{oid} to'lovi tasdiqlandi!</b>\n"
                    f"📦 {o.get('product_name','—')}\n\n"
                    f"🚶 <b>Mahsulotni o'zingiz olib ketasiz.</b>\n"
                    f"🏪 Do'kon: {shop_name}\n"
                    f"🏙 Shahar: {shop_city}\n"
                    f"📱 Do'kon tel: {shop_phone}\n\n"
                    f"Do'kon bilan bog'lanib, mahsulotni olib keting."
                    f"{gift_line_buyer}",
                    parse_mode="HTML"
                )
            except Exception:
                pass

        # Sellerga (ega + yordamchilar) — XARIDOR MA'LUMOTLARI KO'RSATILMAYDI
        seller_msg = (
            f"💳 <b>Buyurtma #{oid} — to'lov tasdiqlandi!</b>\n\n"
            f"📦 {o.get('product_name','—')}\n"
            f"🚚 {dlv_label}\n\n"
            f"🔒 <b>Xaridor o'zi olib ketadi — ma'lumotlari ko'rsatilmaydi.</b>\n"
            f"Xaridor do'kon raqamiga bog'lanib, mahsulotni olib ketadi.\n"
            f"{gift_line_seller}"
            f"💡 <i>Eslatma: bu buyurtma uchun platforma xizmat haqi "
            f"({o.get('commission', int(o.get('total',0)*0.1)):,} so'm — 10%) "
            f"xaridorning oldi-to'lovidan olingan.</i>\n"
            f"Mahsulotni tayyorlab qo'ying. /orders"
        )
        for nid in shop_notify_ids(o["seller_id"]):
            try:
                await bot.send_message(nid, seller_msg, parse_mode="HTML")
            except Exception:
                pass
    else:
        # ── TAKSI/YETKAZIB BERISH: xaridor ma'lumotlari ENDI ochiladi (to'lov tasdiqlangani uchun) ──
        # Xaridorga
        if notify_buyer:
            try:
                await bot.send_message(
                    o["buyer_id"],
                    f"✅ <b>Buyurtma #{oid} to'lovi tasdiqlandi!</b>\n"
                    f"📦 {o.get('product_name','—')}\n"
                    f"{_speed_text(speed)}\n"
                    f"Buyurtmangiz tayyorlanmoqda. 🔄"
                    f"{gift_line_buyer}",
                    parse_mode="HTML"
                )
            except Exception:
                pass

        # Sellerga (ega + yordamchilar) — XARIDOR MA'LUMOTLARI KO'RSATILMAYDI.
        # Endi xaridor bilan to'g'ridan-to'g'ri aloqa yo'q — hamma narsa kurier
        # orqali. Xaridor kontakti faqat kurierga ko'rinadi.
        seller_msg = (
            f"💳 <b>Buyurtma #{oid} — to'lov tasdiqlandi!</b>\n\n"
            f"📦 {o.get('product_name','—')}\n"
            f"🚚 {dlv_label}\n\n"
            f"🔒 <b>Xaridor ma'lumotlari kurierда — sizga ko'rsatilmaydi.</b>\n"
            f"Mahsulotni tayyorlab qo'ying — kurier do'koningizdan olib ketadi.\n"
            f"{gift_line_seller}"
            f"Kurierga berganingizda «🚚 Kurierga berdim» tugmasini bosing.\n\n"
            f"💡 <i>Eslatma: bu buyurtma uchun platforma xizmat haqi "
            f"({o.get('commission', int(o.get('total',0)*0.1)):,} so'm — 10%) "
            f"xaridorning oldi-to'lovidan olingan.</i>\n"
            f"/orders"
        )
        # «🚚 Kurierga berdim» tugmasi — seller mahsulotni kurierga berganda
        # bosadi. Bosilganda buyurtma "Yo'lda" bo'lib, xaridorга "mahsulotingiz
        # yo'lda" xabari ketadi (handler: seller_panel.update_order).
        ship_kb = InlineKeyboardMarkup(inline_keyboard=[[
            InlineKeyboardButton(text="🚚 Kurierga berdim",
                                 callback_data=f"ostatus_{oid}_shipped"),
        ]])
        for nid in shop_notify_ids(o["seller_id"]):
            try:
                await bot.send_message(nid, seller_msg, parse_mode="HTML",
                                       reply_markup=ship_kb)
            except Exception:
                pass

        # ── KURIERLARGA — yangi zakaz (10% to'lov tasdiqlanishi bilanoq) ──
        # Xaridor manzili/telefoni FAQAT shu xabarda ochiladi — sellerga emas.
        couriers = get_couriers()
        courier_ok = 0
        for cid in couriers:
            try:
                await bot.send_message(int(cid), _courier_order_text(oid, o),
                                       parse_mode="HTML",
                                       reply_markup=_courier_done_kb(oid))
                courier_ok += 1
            except Exception:
                # Kurier botga /start bosmagan yoki bloklagan bo'lsa — o'tkazib yuboramiz
                pass
        if not courier_ok:
            # Zakaz birorta ham kurierga yetmadi — ownerga ogohlantirish, aks holda
            # zakaz "havoda qolib" hech kim yetkazmasligi mumkin.
            sabab = ("kurierlar ro'yxati bo'sh" if not couriers
                     else "kurier(lar) botga /start bosmagan yoki botni bloklagan")
            try:
                await bot.send_message(
                    settings.OWNER_ID,
                    f"⚠️ <b>Zakaz #{oid} birorta ham kurierga yetmadi!</b>\n"
                    f"Sabab: {sabab}.\n\n"
                    f"Admin panel → 🚚 Kurierlar → ➕ Kurier qo'shish orqali "
                    f"kurier qo'shing (kurier avval botga /start bosgan bo'lishi shart). "
                    f"Keyin kurier zakazni /kurier buyrug'i bilan ko'ra oladi.",
                    parse_mode="HTML"
                )
            except Exception:
                pass

    # ── AUKSION guruhiga: rasm + buyurtma raqami + soni ──
    # FAQAT shu yerda (to'lov admin tomonidan tasdiqlanganda) yuboriladi —
    # buyurtma yaratilganda emas. Pickup/taksi — ikkalasiga ham tegishli.
    try:
        from app.handlers.common import _send_to_auction
        prod = get_product_by_id(o.get("product_id"))
        photos = product_photos(prod) if prod else []
        await _send_to_auction(photos[0] if photos else None,
                               oid, o.get("quantity", 1))
    except Exception:
        pass

    # Savat (guruh) oqimi xaridorga umumiy xabarda sovg'ani ko'rsatishi uchun.
    return gift_now


@router.callback_query(F.data.startswith("paycfm_"))
async def confirm_payment(call: CallbackQuery):
    if not is_owner(call.from_user.id):
        await call.answer("Ruxsat yo'q.", show_alert=True); return
    # Yetkazib berish vaqtini tugma suffiksidan o'qiymiz: paycfm_{oid}_today|_24h
    data = call.data
    if data.endswith("_24h"):
        speed = "24h"; data = data[:-4]
    elif data.endswith("_today"):
        speed = "today"; data = data[:-6]
    else:
        speed = "today"   # eski xabarlar uchun (vaqt belgilanmagan) — bugun
    # Tugma spinnerini DARHOL to'xtatamiz — "qotib qolish" oldini oladi
    await call.answer("✅ To'lov tasdiqlandi!")

    oid = int(data.split("_")[1])
    o = get_order_by_id(oid)
    if not o:
        try: await call.message.answer("❌ Buyurtma topilmadi.")
        except Exception: pass
        return

    speed_label = "24 soat ichida" if speed == "24h" else "BUGUN"
    try:
        _log(call.from_user, "To'lov tasdiqlandi",
             f"Buyurtma #{oid} — yetkazish: {speed_label}")
    except Exception:
        pass
    try:
        await call.message.edit_caption(
            caption=(call.message.caption or "")
                    + f"\n\n✅ TO'LOV TASDIQLANDI — 🚚 {speed_label}",
        )
    except Exception:
        pass

    await _confirm_single_order(o, speed=speed)


@router.callback_query(F.data.startswith("paycfmg_"))
async def confirm_payment_group(call: CallbackQuery):
    """Savat buyurtmasi (guruh) — barcha mahsulotlar to'lovini bir bosishda tasdiqlaydi."""
    if not is_owner(call.from_user.id):
        await call.answer("Ruxsat yo'q.", show_alert=True); return
    # Yetkazib berish vaqti suffiksi: paycfmg_{group_id}_today|_24h
    # (group_id ichida ham "_" bor, shuning uchun suffiks bo'yicha ajratamiz.)
    data = call.data
    if data.endswith("_24h"):
        speed = "24h"; data = data[:-4]
    elif data.endswith("_today"):
        speed = "today"; data = data[:-6]
    else:
        speed = "today"
    await call.answer("✅ To'lov tasdiqlandi!")

    group_id = data.split("_", 1)[1]
    orders = get_orders_by_group(group_id)
    if not orders:
        try: await call.message.answer("❌ Buyurtmalar topilmadi.")
        except Exception: pass
        return

    try:
        _log(call.from_user, "Savat to'lovi tasdiqlandi",
             f"Guruh {group_id} — {len(orders)} ta buyurtma")
    except Exception:
        pass
    speed_label = "24 soat ichida" if speed == "24h" else "BUGUN"
    try:
        await call.message.edit_caption(
            caption=(call.message.caption or "")
                    + f"\n\n✅ TO'LOV TASDIQLANDI — 🚚 {speed_label}",
        )
    except Exception:
        pass

    # Har bir buyurtmani alohida tasdiqlaymiz (seller/kurier xabarlari), lekin
    # xaridorга bitta umumiy xabar yuboramiz. Sovg'a bitta profilga bir marta —
    # guruhda ham faqat birinchi buyurtma sovg'a oladi.
    gift_any = False
    for o in orders:
        if await _confirm_single_order(o, notify_buyer=False, speed=speed):
            gift_any = True

    items_txt = "".join(
        f"• {o.get('product_name','—')} — {o.get('quantity',1)} dona\n"
        for o in orders
    )
    gift_line = (
        "\n🎁 Buyurtmangizga <b>kafolatlangan sovg'a</b> qo'shib yuboriladi!"
        if gift_any else ""
    )
    try:
        from app.bot.bot import bot
        await bot.send_message(
            orders[0]["buyer_id"],
            f"✅ <b>Buyurtmangiz to'lovi tasdiqlandi!</b>  ({len(orders)} ta mahsulot)\n\n"
            f"{items_txt}\n"
            f"{_speed_text(speed)}\n"
            f"Buyurtmangiz tayyorlanmoqda. 🔄"
            f"{gift_line}",
            parse_mode="HTML",
        )
    except Exception:
        pass


@router.callback_query(F.data.startswith("payrejg_"))
async def reject_payment_group(call: CallbackQuery):
    """Savat buyurtmasi (guruh) chekini rad etadi."""
    if not is_owner(call.from_user.id): return
    group_id = call.data.split("_", 1)[1]
    orders = get_orders_by_group(group_id)
    if not orders:
        await call.answer("Buyurtmalar topilmadi.", show_alert=True); return
    await call.answer("❌ Rad etildi.")
    try:
        await call.message.edit_caption(
            caption=(call.message.caption or "") + "\n\n❌ CHEK RAD ETILDI",
        )
    except Exception:
        pass
    kb = InlineKeyboardMarkup(inline_keyboard=[[
        InlineKeyboardButton(text="🧾 Chekni qayta yuborish",
                             callback_data=f"gsendrcpt_{group_id}")
    ]])
    try:
        from app.bot.bot import bot
        await bot.send_message(
            orders[0]["buyer_id"],
            f"❌ <b>Savat buyurtmangiz cheki rad etildi.</b>\n"
            f"Iltimos, to'g'ri to'lov chekini qayta yuboring:",
            parse_mode="HTML", reply_markup=kb,
        )
    except Exception:
        pass


@router.callback_query(F.data.startswith("payrej_"))
async def reject_payment(call: CallbackQuery):
    if not is_owner(call.from_user.id): return
    oid = int(call.data.split("_")[1])
    o = get_order_by_id(oid)
    if not o:
        await call.answer("Buyurtma topilmadi.", show_alert=True); return
    await call.answer("❌ Rad etildi.")
    try:
        await call.message.edit_caption(
            caption=(call.message.caption or "") + "\n\n❌ CHEK RAD ETILDI",
        )
    except Exception:
        pass
    try:
        from app.bot.bot import bot
        await bot.send_message(
            o["buyer_id"],
            f"❌ <b>Buyurtma #{oid} cheki rad etildi.</b>\n"
            f"Iltimos, to'g'ri to'lov chekini qayta yuboring:\n"
            f"📦 Buyurtmalarim → 🧾 chek yuborish",
            parse_mode="HTML"
        )
    except Exception:
        pass


# ═══════════════════════════════════════════════════════════════════════════
#  PROMO-KODLAR (chegirma)
# ═══════════════════════════════════════════════════════════════════════════
def _promos_view():
    promos = get_promos()
    if not promos:
        text = "🎁 <b>Promo-kodlar</b>\n\nHozircha promo-kod yo'q."
    else:
        text = f"🎁 <b>Promo-kodlar</b> — {len(promos)} ta\n\n"
        for code, p in promos.items():
            limit = int(p.get("limit", 0))
            used  = int(p.get("used", 0))
            limit_txt = f"{used}/{limit}" if limit else f"{used}/∞"
            text += f"• <code>{code}</code> — −{p.get('percent',0)}%  (ishlatilgan: {limit_txt})\n"
        text += "\nO'chirish uchun tugmani bosing."
    rows = [[InlineKeyboardButton(text=f"🗑 {code}", callback_data=f"delpromo_{code}")]
            for code in get_promos()]
    rows.append([InlineKeyboardButton(text="➕ Promo-kod qo'shish", callback_data="addpromo")])
    rows.append([InlineKeyboardButton(text="🔙 Orqaga", callback_data="admin_back")])
    return text, InlineKeyboardMarkup(inline_keyboard=rows)


@router.callback_query(F.data == "admin_promos")
async def admin_promos(call: CallbackQuery):
    if not is_owner(call.from_user.id): return
    await _ack(call)
    text, kb = _promos_view()
    await _admin_nav(call, text, kb)


@router.callback_query(F.data == "addpromo")
async def add_promo_start(call: CallbackQuery, state: FSMContext):
    if not is_owner(call.from_user.id): return
    await _ack(call)
    await state.set_state(AdminPromo.code)
    await call.message.answer(
        "🎁 Yangi promo-kod nomini kiriting (masalan: SAYL10):"
    )


@router.message(AdminPromo.code)
async def add_promo_code(message: Message, state: FSMContext):
    code = "".join((message.text or "").split()).upper()
    if not code or len(code) > 24:
        await message.answer("❌ Kodni to'g'ri kiriting (probelsiz, 24 belgigacha):"); return
    if get_promo(code):
        await message.answer("❌ Bu kod allaqachon bor. Boshqa nom kiriting:"); return
    await state.update_data(code=code)
    await state.set_state(AdminPromo.percent)
    await message.answer("💯 Chegirma foizini kiriting (1–90):")


@router.message(AdminPromo.percent)
async def add_promo_percent(message: Message, state: FSMContext):
    pct = to_int(message.text, -1)
    if pct < 1 or pct > 90:
        await message.answer("❌ Foizni 1 dan 90 gacha kiriting:"); return
    await state.update_data(percent=pct)
    await state.set_state(AdminPromo.limit)
    await message.answer(
        "🔢 Necha marta ishlatish mumkin? Sonini kiriting,\n"
        "yoki cheksiz uchun /skip yozing:"
    )


@router.message(AdminPromo.limit)
async def add_promo_limit(message: Message, state: FSMContext):
    txt = (message.text or "").strip()
    if txt == "/skip":
        limit = 0
    else:
        limit = to_int(txt, -1)
        if limit < 0:
            await message.answer("❌ Sonni to'g'ri kiriting yoki /skip yozing:"); return
    data = await state.get_data()
    ok = add_promo(data["code"], data["percent"], limit)
    await state.clear()
    if not ok:
        await message.answer("❌ Saqlab bo'lmadi (kod allaqachon bor bo'lishi mumkin).",
                             reply_markup=main_menu)
        return
    _log(message.from_user, "Promo-kod qo'shildi",
         f"{data['code']} −{data['percent']}% limit={limit or '∞'}")
    limit_txt = str(limit) if limit else "cheksiz"
    await message.answer(
        f"✅ Promo-kod qo'shildi!\n\n"
        f"🎁 <code>{data['code']}</code>\n"
        f"💯 Chegirma: −{data['percent']}%\n"
        f"🔢 Limit: {limit_txt}",
        parse_mode="HTML", reply_markup=main_menu
    )


@router.callback_query(F.data.startswith("delpromo_"))
async def del_promo_cb(call: CallbackQuery):
    if not is_owner(call.from_user.id): return
    code = call.data.split("_", 1)[1]
    if delete_promo(code):
        await call.answer(f"🗑 {code} o'chirildi")
        _log(call.from_user, "Promo-kod o'chirildi", code)
    else:
        await call.answer("Topilmadi.")
    text, kb = _promos_view()
    await _admin_nav(call, text, kb)


# ═══════════════════════════════════════════════════════════════════════════
#  HISOBOTLAR (Excel / Word eksportlari bitta joyda)
# ═══════════════════════════════════════════════════════════════════════════
@router.callback_query(F.data == "admin_reports")
async def admin_reports(call: CallbackQuery):
    if not is_owner(call.from_user.id): return
    await _ack(call)
    kb = InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="📈 Excel hisobot",    callback_data="admin_excel_menu")],
        [InlineKeyboardButton(text="📄 Sellerlar (Word)", callback_data="admin_sellers_word")],
        [InlineKeyboardButton(text="📜 Jurnal (Word)",    callback_data="admin_log")],
        [InlineKeyboardButton(text="🔙 Orqaga",           callback_data="admin_back")],
    ])
    await _admin_nav(call, "📑 <b>Hisobotlar</b>\n\nKerakli hisobotni tanlang:", kb)


# ─── Sellerlar (ko'rish + tahrirlash + o'chirish) ────────────────────────────
@router.callback_query(F.data == "admin_sellers")
async def show_sellers(call: CallbackQuery):
    if not is_owner(call.from_user.id): return
    await _ack(call)
    sellers = get_sellers()
    if not sellers:
        await _admin_nav(call, "Hozircha seller yo'q.", InlineKeyboardMarkup(
            inline_keyboard=[[InlineKeyboardButton(text="🔙 Orqaga", callback_data="admin_back")]]
        ))
        return
    items = list(sellers.items())
    await _admin_nav(call, _sellers_text(items), _sellers_keyboard(items))


# ─── Seller qidirish ─────────────────────────────────────────────────────────
@router.callback_query(F.data == "seller_search")
async def seller_search_start(call: CallbackQuery, state: FSMContext):
    if not is_owner(call.from_user.id): return
    await state.set_state(AdminSellerSearch.query)
    await call.message.answer("🔍 Do'kon nomi yoki seller ismini kiriting:")
    await call.answer()


@router.message(AdminSellerSearch.query)
async def seller_search_do(message: Message, state: FSMContext):
    if not is_owner(message.from_user.id):
        await state.clear(); return
    await state.clear()
    q = (message.text or "").strip().lower()
    items = [
        (uid, s) for uid, s in get_sellers().items()
        if q in s.get("shop_name", "").lower() or q in s.get("full_name", "").lower()
    ]
    if not items:
        await message.answer(
            f"😔 '{message.text}' bo'yicha seller topilmadi.",
            reply_markup=InlineKeyboardMarkup(inline_keyboard=[
                [InlineKeyboardButton(text="🏪 Barcha sellerlar", callback_data="admin_sellers")]
            ])
        )
        return
    await message.answer(_sellers_text(items), parse_mode="HTML",
                         reply_markup=_sellers_keyboard(items))


@router.callback_query(F.data.startswith("aseller_"))
async def seller_detail(call: CallbackQuery):
    if not is_owner(call.from_user.id): return
    uid = call.data.split("_")[1]
    s = get_sellers().get(uid)
    if not s:
        await call.answer("Seller topilmadi."); return
    await _ack(call)
    from app.storage import get_seller_rating, get_seller_products
    rating, cnt = get_seller_rating(int(uid))
    stars = "⭐" * int(rating) if rating else "—"
    prods = get_seller_products(int(uid))
    text = (
        f"🏪 <b>{s['shop_name']}</b>\n\n"
        f"👤 {s['full_name']}\n"
        f"🏙 Shahar: {s.get('city','—')}\n"
        f"📱 {s['phone']}\n"
        f"💳 {s.get('card_number','—')}\n"
        f"🆔 ID: {uid}\n"
        f"📦 Mahsulotlar: {len(prods)} ta\n"
        f"⭐ Reyting: {stars} {rating} ({cnt} ta baho)"
    )
    kb = InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="✉️ Sellerga yozish (lichka)", callback_data=f"msg_seller_{uid}")],
        [InlineKeyboardButton(text="✏️ Do'kon nomini o'zgartirish", callback_data=f"edit_shop_name_{uid}")],
        [InlineKeyboardButton(text="✏️ Karta raqamini o'zgartirish", callback_data=f"edit_shop_card_{uid}")],
        [InlineKeyboardButton(text="✏️ Telefon raqamini o'zgartirish", callback_data=f"edit_shop_phone_{uid}")],
        [InlineKeyboardButton(text="🗑 Sellerni o'chirish", callback_data=f"del_seller_{uid}")],
        [InlineKeyboardButton(text="🔙 Orqaga", callback_data="admin_sellers")],
    ])
    await _admin_nav(call, text, kb)


@router.callback_query(F.data.startswith("edit_shop_"))
async def edit_shop_start(call: CallbackQuery, state: FSMContext):
    if not is_owner(call.from_user.id): return
    parts = call.data.split("_")          # edit_shop_name_123  → ['edit','shop','name','123']
    field = parts[2]                       # name | card | phone
    uid   = parts[3]
    field_labels = {"name": "Do'kon nomi", "card": "Karta raqami", "phone": "Telefon"}
    await state.set_state(AdminEditShop.waiting_value)
    await state.update_data(field=field, uid=uid)
    await call.message.answer(f"✏️ Yangi <b>{field_labels.get(field,'qiymat')}</b>ni kiriting:", parse_mode="HTML")
    await call.answer()


@router.message(AdminEditShop.waiting_value)
async def edit_shop_save(message: Message, state: FSMContext):
    if not is_owner(message.from_user.id):
        await state.clear(); return
    data = await state.get_data()
    field = data["field"]
    uid   = int(data["uid"])
    mapping = {"name": "shop_name", "card": "card_number", "phone": "phone"}
    update_seller(uid, {mapping[field]: message.text})
    await state.clear()
    await message.answer(f"✅ Seller ma'lumoti yangilandi!")


# ─── Sellerga to'g'ridan to'g'ri xabar yuborish (lichkasiga) ─────────────────
@router.callback_query(F.data.startswith("msg_seller_"))
async def msg_seller_start(call: CallbackQuery, state: FSMContext):
    if not is_owner(call.from_user.id): return
    uid = call.data.split("_")[2]
    s = get_sellers().get(uid)
    if not s:
        await call.answer("Seller topilmadi."); return
    u = get_user(int(uid)) or {}
    uname = f"@{u['username']}" if u.get("username") else "—"
    await state.set_state(AdminMsgSeller.text)
    await state.update_data(target_uid=int(uid), shop_name=s["shop_name"])
    await call.message.answer(
        f"✉️ <b>{s['shop_name']}</b> ({s['full_name']}) selleriga xabar yozing.\n"
        f"👤 Username: {uname}\n\n"
        f"Matn, rasm yoki istalgan xabarni yuboring — bot uni sellerning "
        f"lichkasiga yetkazadi.\n"
        f"Bekor qilish: /cancel",
        parse_mode="HTML"
    )
    await call.answer()


@router.message(AdminMsgSeller.text)
async def msg_seller_send(message: Message, state: FSMContext):
    if not is_owner(message.from_user.id):
        await state.clear(); return
    if (message.text or "").startswith("/"):
        await state.clear()
        await message.answer("⛔️ Xabar yuborish bekor qilindi.",
                             reply_markup=admin_menu_kb(message.from_user.id))
        return
    data = await state.get_data()
    uid  = data["target_uid"]
    shop = data.get("shop_name", "—")
    await state.clear()
    from app.bot.bot import bot
    try:
        await bot.send_message(uid, "📩 <b>Admindan xabar:</b>", parse_mode="HTML")
        await message.send_copy(chat_id=uid)
    except Exception:
        await message.answer(
            "❌ Xabar yetib bormadi. Seller botni bloklagan yoki "
            "hali /start bosmagan bo'lishi mumkin."
        )
        return
    _log(message.from_user, "Sellerga xabar yuborildi", f"{shop} (ID:{uid})")
    await message.answer(
        f"✅ Xabar <b>{shop}</b> selleriga yuborildi!", parse_mode="HTML",
        reply_markup=InlineKeyboardMarkup(inline_keyboard=[
            [InlineKeyboardButton(text="✉️ Yana yozish", callback_data=f"msg_seller_{uid}")],
            [InlineKeyboardButton(text="🔙 Sellerlar", callback_data="admin_sellers")],
        ])
    )


# ─── 🔄 Restart: hammaga yangi menyuni majburan yuborish ─────────────────────
# Foydalanuvchilar /start bosmasdan yoki tarixini tozalamasdan ham darhol
# yangi klaviaturani oladi.
@router.callback_query(F.data == "admin_restart_menu")
async def admin_restart_menu_confirm(call: CallbackQuery):
    if not is_owner(call.from_user.id): return
    targets = set(get_users().keys()) | set(get_sellers().keys())
    if not targets:
        await call.answer("Foydalanuvchilar yo'q.", show_alert=True); return
    await call.message.answer(
        f"🔄 <b>Restart — tarixni tozalash</b>\n\n"
        f"Bu amal <b>{len(targets)}</b> ta foydalanuvchi uchun bajariladi:\n"
        f"🧹 Barcha <b>buyurtmalar</b>, <b>sharhlar (reytinglar)</b> va "
        f"<b>xaridor profillari</b> o'chiriladi — bot bo'sh holatga qaytadi.\n"
        f"❤️ <b>Istaklar (sevimlilar) saqlanadi.</b>\n"
        f"🏪 Do'konlar va mahsulotlar saqlanadi.\n"
        f"👋 Har bir foydalanuvchiga avtomatik yangi /start ekrani yuboriladi.\n\n"
        f"⚠️ <b>Bu amalni ortga qaytarib bo'lmaydi.</b> Tasdiqlaysizmi?",
        parse_mode="HTML",
        reply_markup=InlineKeyboardMarkup(inline_keyboard=[
            [InlineKeyboardButton(text="✅ Ha, tozalab yuborish", callback_data="admin_restart_go")],
            [InlineKeyboardButton(text="🔙 Orqaga", callback_data="admin_back")],
        ])
    )
    await _ack(call)


async def _send_start_screen(bot, uid: int):
    """Foydalanuvchiga /start ekranini yuboradi (restartdan keyin avtomatik).
    Seller bo'lsa — seller menyusi; xaridor bo'lsa — salomlashish + "Seller
    bo'lish" tugmasi + pastki menyu (xuddi /start bosgandek)."""
    from app.handlers.common import TOP_BANNER, SELLER_INVITE_BANNER
    from app.keyboards.seller import menu_for, seller_main_menu
    from app.storage import is_shop_member
    if is_shop_member(uid):
        await bot.send_message(
            uid, "🔄 Bot yangilandi. Quyidagi menyudan foydalaning:",
            reply_markup=seller_main_menu,
        )
    else:
        kb = InlineKeyboardMarkup(inline_keyboard=[
            [InlineKeyboardButton(text="🏪 Seller bo'lish", callback_data="become_seller")],
            [InlineKeyboardButton(text="💬 Adminga yozish",
                                  url=f"https://t.me/{settings.ADMIN_USERNAME}")],
        ])
        await bot.send_message(
            uid,
            f"{TOP_BANNER}\n\n"
            "👋 Salom!\n<b>Proman Market</b> botiga xush kelibsiz!\n\n"
            f"{SELLER_INVITE_BANNER}",
            parse_mode="HTML", reply_markup=kb,
        )
        await bot.send_message(
            uid, "👇 Yoki quyidagi menyudan tanlang:", reply_markup=menu_for(uid),
        )
    set_user_field(uid, "menu_ver", MENU_VERSION)


# ── Restart: ikki bosqichli tasdiq ──
# 1-bosqich: tasdiq ekrani (tugma) → 2-bosqich: "TOZALA" so'zini yozib tasdiqlash.
# Bitta tugma bilan adashib o'chirib yuborilmasligi uchun yozma tasdiq talab qilamiz.
@router.callback_query(F.data == "admin_restart_go")
async def admin_restart_ask_word(call: CallbackQuery, state: FSMContext):
    if not is_owner(call.from_user.id): return
    await _ack(call)
    await state.set_state(AdminRestartConfirm.word)
    await call.message.answer(
        "⚠️ <b>Oxirgi tasdiq.</b>\n"
        "Bu amal barcha buyurtma, sharh va xaridor profillarini o'chiradi.\n\n"
        "Davom etish uchun katta harflar bilan <b>TOZALA</b> deb yozing.\n"
        "Bekor qilish: /cancel",
        parse_mode="HTML"
    )


@router.message(AdminRestartConfirm.word)
async def admin_restart_confirm_word(message: Message, state: FSMContext):
    if not is_owner(message.from_user.id):
        await state.clear(); return
    txt = (message.text or "").strip()
    await state.clear()
    if txt != "TOZALA":
        await message.answer(
            "❌ Restart bekor qilindi (tasdiqlash so'zi noto'g'ri).",
            reply_markup=main_menu
        )
        return
    await _run_restart(message.from_user, message)


async def _run_restart(actor, message):
    from app.bot.bot import bot
    from app.storage import reset_buyer_data, pop_all_view_msgs

    # 1) Foydalanuvchilar ro'yxatini TOZALASHDAN OLDIN yig'ib olamiz
    #    (reset_buyer_data users.json'ni bo'shatadi — keyin ro'yxat yo'qoladi).
    targets = sorted(set(get_users().keys()) | set(get_sellers().keys()))
    status = await message.answer(f"🔄 Tozalanmoqda...  0/{len(targets)}")

    # 2) Chatlardagi ko'rsatilgan mahsulot kartochkalarini o'chiramiz
    #    ("ko'rgan narsalari" — chat imkon qadar tozalanadi).
    for chat_id, ids in pop_all_view_msgs().items():
        await asyncio.gather(
            *(bot.delete_message(chat_id, mid) for mid in ids),
            return_exceptions=True,
        )

    # 3) Xaridor tarixini tozalaymiz (istaklar/do'konlar/mahsulotlar saqlanadi)
    counts = reset_buyer_data()

    # 4) Har bir foydalanuvchiga avtomatik yangi /start ekrani yuboramiz
    sent, failed = 0, 0
    for i, uid in enumerate(targets, 1):
        try:
            await _send_start_screen(bot, int(uid))
            sent += 1
        except Exception:
            failed += 1
        # Telegram flood limitiga tushib qolmaslik uchun kichik pauza
        await asyncio.sleep(0.1)
        if i % 25 == 0:
            try:
                await status.edit_text(f"🔄 Tozalanmoqda...  {i}/{len(targets)}")
            except Exception:
                pass

    _log(actor, "Restart: tarix tozalandi + start",
         f"buyurtma:{counts['orders']}, sharh:{counts['reviews']}, "
         f"profil:{counts['users']}; yetdi:{sent}, yetmadi:{failed}")
    result = (
        f"🔄 <b>Restart yakunlandi!</b>\n"
        f"🧹 O'chirildi: <b>{counts['orders']}</b> buyurtma, "
        f"<b>{counts['reviews']}</b> sharh, <b>{counts['users']}</b> profil\n"
        f"❤️ Istaklar saqlandi.\n"
        f"✅ Yangi start yetib bordi: <b>{sent}</b> ta\n"
    )
    if failed:
        result += (
            f"❌ Yetib bormadi: <b>{failed}</b> ta "
            f"(botni bloklagan yoki hali /start bosmagan)\n"
        )
    try:
        await status.edit_text(result, parse_mode="HTML")
    except Exception:
        await message.answer(result, parse_mode="HTML")


# ─── Ommaviy xabar / e'lon (sellerlar, xaridorlar yoki hammaga) ──────────────
def _broadcast_targets(target: str) -> list[int]:
    """Ommaviy xabar uchun qabul qiluvchilar id ro'yxati.
    target: "sellers" | "buyers" | "all".
    - sellers : barcha sellerlar
    - buyers  : seller bo'lmagan oddiy xaridorlar (botga /start bosganlar)
    - all     : sellerlar + xaridorlar (hamma)
    """
    seller_ids = {int(uid) for uid in get_sellers()}
    user_ids = {int(uid) for uid in get_users()}
    if target == "sellers":
        ids = seller_ids
    elif target == "buyers":
        ids = user_ids - seller_ids
    else:  # all
        ids = user_ids | seller_ids
    return sorted(ids)


_BROADCAST_LABELS = {
    "sellers": "🏪 Sellerlar",
    "buyers":  "🛍 Xaridorlar",
    "all":     "📢 Hamma (seller + xaridor)",
}


@router.callback_query(F.data == "admin_broadcast")
async def admin_broadcast_menu(call: CallbackQuery):
    if not is_owner(call.from_user.id): return
    await _ack(call)
    n_sellers = len(_broadcast_targets("sellers"))
    n_buyers  = len(_broadcast_targets("buyers"))
    n_all     = len(_broadcast_targets("all"))
    text = (
        "📢 <b>Ommaviy xabar / e'lon</b>\n\n"
        "Kimga yubormoqchisiz? Tanlang — keyin matn, rasm yoki "
        "istalgan xabarni yuboring, bot uni har bir kishining lichkasiga "
        "yetkazadi.\n\n"
        f"🏪 Sellerlar: <b>{n_sellers}</b> ta\n"
        f"🛍 Xaridorlar: <b>{n_buyers}</b> ta\n"
        f"📢 Hammasi: <b>{n_all}</b> ta"
    )
    kb = InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text=f"🏪 Sellerlarga ({n_sellers})", callback_data="bcast_sellers")],
        [InlineKeyboardButton(text=f"🛍 Xaridorlarga ({n_buyers})", callback_data="bcast_buyers")],
        [InlineKeyboardButton(text=f"📢 Hammaga ({n_all})", callback_data="bcast_all")],
        [InlineKeyboardButton(text="🔙 Orqaga", callback_data="admin_back")],
    ])
    await _admin_nav(call, text, kb)


@router.callback_query(F.data.startswith("bcast_"))
async def admin_broadcast_start(call: CallbackQuery, state: FSMContext):
    if not is_owner(call.from_user.id): return
    target = call.data.split("_", 1)[1]   # bcast_sellers → "sellers"
    if target not in _BROADCAST_LABELS:
        await call.answer("Noma'lum tanlov."); return
    recipients = _broadcast_targets(target)
    if not recipients:
        await call.answer("Bu guruhda hech kim yo'q.", show_alert=True); return
    await state.set_state(AdminBroadcast.text)
    await state.update_data(target=target)
    await call.message.answer(
        f"📢 <b>{_BROADCAST_LABELS[target]}</b>ga xabar  (jami {len(recipients)} ta)\n\n"
        f"Endi matn, rasm, video yoki istalgan xabarni yuboring — bot uni "
        f"har bir kishining lichkasiga yetkazadi.\n"
        f"Bekor qilish: /cancel",
        parse_mode="HTML"
    )
    await call.answer()


@router.message(AdminBroadcast.text)
async def admin_broadcast_send(message: Message, state: FSMContext):
    if not is_owner(message.from_user.id):
        await state.clear(); return
    if (message.text or "").startswith("/"):
        await state.clear()
        await message.answer("⛔️ Xabar yuborish bekor qilindi.",
                             reply_markup=admin_menu_kb(message.from_user.id))
        return
    data = await state.get_data()
    target = data.get("target", "sellers")
    await state.clear()
    recipients = _broadcast_targets(target)
    from app.bot.bot import bot
    status = await message.answer(f"📤 Yuborilmoqda...  0/{len(recipients)}")
    sent, failed = 0, 0
    for i, uid in enumerate(recipients, 1):
        try:
            await bot.send_message(int(uid), "📩 <b>Admindan xabar:</b>", parse_mode="HTML")
            await message.send_copy(chat_id=int(uid))
            sent += 1
        except Exception:
            # Botni bloklagan yoki hali /start bosmagan — o'tkazib yuboramiz
            failed += 1
        # Telegram flood limitiga tushib qolmaslik uchun kichik pauza
        await asyncio.sleep(0.1)
        if i % 25 == 0:
            try:
                await status.edit_text(f"📤 Yuborilmoqda...  {i}/{len(recipients)}")
            except Exception:
                pass
    _log(message.from_user, f"Ommaviy xabar ({_BROADCAST_LABELS[target]})",
         f"yetdi: {sent}, yetmadi: {failed}")
    result = (
        f"📢 <b>Yuborish yakunlandi!</b>\n"
        f"🎯 Maqsad: {_BROADCAST_LABELS[target]}\n"
        f"✅ Yetib bordi: <b>{sent}</b> ta\n"
    )
    if failed:
        result += (
            f"❌ Yetib bormadi: <b>{failed}</b> ta "
            f"(botni bloklagan yoki /start bosmagan)\n"
        )
    await message.answer(result, parse_mode="HTML",
                         reply_markup=admin_menu_kb(message.from_user.id))


@router.callback_query(F.data.startswith("del_seller_"))
async def del_seller_cb(call: CallbackQuery):
    if not is_owner(call.from_user.id): return
    uid = int(call.data.split("_")[2])
    sellers = get_sellers()
    name = sellers.get(str(uid), {}).get("shop_name", "—")
    delete_seller(uid)
    _log(call.from_user, "Seller o'chirildi", f"{name} (ID:{uid})")
    await call.message.edit_text(f"🗑 {name} selleri o'chirildi.")
    await call.answer("O'chirildi")


# ─── Mahsulotlar (admin) ─────────────────────────────────────────────────────
# Bitta sahifadagi mahsulotlar soni. Telegram inline-klaviaturasi juda ko'p
# tugmani qabul qilmaydi (mahsulot soni 100 dan oshganda "Mahsulotlar" tugmasi
# umuman ochilmasdi) — shuning uchun sahifalab ko'rsatamiz.
_PROD_PER_PAGE = 30


def _products_page(page: int):
    """(matn, klaviatura) — mahsulotlarning `page`-sahifasi. Sahifalash tugmalari
    (◀️ ▶️) bilan."""
    products = get_all_products()
    total = len(products)
    pages = max(1, (total + _PROD_PER_PAGE - 1) // _PROD_PER_PAGE)
    page = max(0, min(page, pages - 1))
    start = page * _PROD_PER_PAGE
    rows = [
        [InlineKeyboardButton(text=f"📦 {p.get('name','—')} — {p.get('price',0):,} so'm",
                              callback_data=f"aprod_{p['id']}")]
        for p in products[start:start + _PROD_PER_PAGE]
    ]
    nav = []
    if page > 0:
        nav.append(InlineKeyboardButton(text="◀️ Oldingi", callback_data=f"aprodpg_{page-1}"))
    if page < pages - 1:
        nav.append(InlineKeyboardButton(text="Keyingi ▶️", callback_data=f"aprodpg_{page+1}"))
    if nav:
        rows.append(nav)
    rows.append([InlineKeyboardButton(text="🔙 Orqaga", callback_data="admin_back")])
    text = f"📦 <b>Barcha mahsulotlar</b> — {total} ta  (sahifa {page+1}/{pages}):"
    return text, InlineKeyboardMarkup(inline_keyboard=rows)


@router.callback_query(F.data == "admin_products")
async def admin_products(call: CallbackQuery):
    if not is_owner(call.from_user.id): return
    await _ack(call)
    if not get_all_products():
        await _admin_nav(call, "Mahsulot yo'q.", InlineKeyboardMarkup(
            inline_keyboard=[[InlineKeyboardButton(text="🔙 Orqaga", callback_data="admin_back")]]
        ))
        return
    text, kb = _products_page(0)
    await _admin_nav(call, text, kb)


@router.callback_query(F.data.startswith("aprodpg_"))
async def admin_products_page(call: CallbackQuery):
    if not is_owner(call.from_user.id): return
    await _ack(call)
    page = int(call.data.split("_")[1])
    text, kb = _products_page(page)
    await _admin_nav(call, text, kb)


def _admin_product_card(p: dict, pid: int):
    """Admin uchun mahsulot kartochkasi matni va tugmalari."""
    colors_line = f"\n🎨 Ranglar: {', '.join(p['colors'])}" if p.get("colors") else ""
    photo_line = f"\n🖼 Rasm: {len(product_photos(p))} ta"
    video_line = "\n🎬 Video: bor" if p.get("video") else "\n🎬 Video: yo'q"
    text = (
        f"📦 <b>{p['name']}</b>\n"
        f"🏪 {p.get('shop_name','—')}\n"
        f"📝 {p.get('description','—')}\n"
        f"💰 {p['price']:,} so'm"
        f"{colors_line}"
        f"{photo_line}{video_line}"
    )
    vid_btn = "🎬 Videoni o'zgartirish" if p.get("video") else "🎬 Video qo'shish"
    kb = InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="✏️ Nomini o'zgartirish",    callback_data=f"eprod_name_{pid}")],
        [InlineKeyboardButton(text="✏️ Narxini o'zgartirish",   callback_data=f"eprod_price_{pid}")],
        [InlineKeyboardButton(text="✏️ Tavsifini o'zgartirish", callback_data=f"eprod_desc_{pid}")],
        [InlineKeyboardButton(text="🎨 Ranglarini o'zgartirish", callback_data=f"eprod_colors_{pid}")],
        [InlineKeyboardButton(text="🖼 Rasmlarni o'zgartirish",  callback_data=f"admphoto_{pid}")],
        [InlineKeyboardButton(text=vid_btn,                      callback_data=f"admvideo_{pid}")],
        [InlineKeyboardButton(text="🗑 O'chirish",               callback_data=f"dprod_{pid}")],
        [InlineKeyboardButton(text="🔙 Orqaga",                  callback_data="admin_products")],
    ])
    return text, kb


async def _admin_send_product(message, p: dict, pid: int):
    """Mahsulot kartochkasini rasm bilan yuboradi (rasm bo'lsa caption sifatida)."""
    text, kb = _admin_product_card(p, pid)
    photos = product_photos(p)
    if photos:
        try:
            await message.answer_photo(photos[0], caption=text, parse_mode="HTML", reply_markup=kb)
            return
        except Exception:
            pass
    await message.answer(text, parse_mode="HTML", reply_markup=kb)


@router.callback_query(F.data.startswith("aprod_"))
async def admin_product_detail(call: CallbackQuery):
    if not is_owner(call.from_user.id): return
    pid = int(call.data.split("_")[1])
    p = get_product_by_id(pid)
    if not p:
        await call.answer("Topilmadi.", show_alert=True); return
    # Ro'yxat (matn) xabarini o'chirib, rasmli kartochkani yuboramiz.
    try:
        await call.message.delete()
    except Exception:
        pass
    await _admin_send_product(call.message, p, pid)
    await call.answer()


@router.callback_query(F.data.startswith("eprod_"))
async def edit_product_start(call: CallbackQuery, state: FSMContext):
    if not is_owner(call.from_user.id): return
    parts = call.data.split("_")   # eprod_name_5
    field = parts[1]
    pid   = int(parts[2])
    labels = {
        "name":   "Nom",
        "price":  "Narx (faqat raqam)",
        "desc":   "Tavsif",
        "colors": "Ranglar (vergul bilan, masalan: Qizil, Ko'k) yoki bo'sh qoldiring",
    }
    await state.set_state(AdminEditProduct.waiting_value)
    await state.update_data(field=field, pid=pid)
    await call.message.answer(f"✏️ Yangi <b>{labels.get(field,'qiymat')}</b>ni kiriting:", parse_mode="HTML")
    await call.answer()


@router.message(AdminEditProduct.waiting_value)
async def edit_product_save(message: Message, state: FSMContext):
    if not is_owner(message.from_user.id):
        await state.clear(); return
    data  = await state.get_data()
    field = data["field"]
    pid   = data["pid"]
    txt = (message.text or "").strip()
    if field != "colors" and not txt:
        await message.answer("❌ Matn ko'rinishida kiriting:"); return
    if field == "price" and not txt.isdigit():
        await message.answer("❌ Narx faqat raqam bo'lishi kerak:"); return
    mapping = {"name": "name", "price": "price", "desc": "description"}
    value = int(txt) if field == "price" else txt
    if field == "colors":
        colors = [c.strip() for c in txt.split(",") if c.strip()] if txt else []
        update_product(pid, {"colors": colors})
    else:
        update_product(pid, {mapping[field]: value})
    await state.clear()
    p = get_product_by_id(pid)
    name = p["name"] if p else f"#{pid}"
    await message.answer(f"✅ <b>{name}</b> yangilandi!", parse_mode="HTML")
    # Yangilangan mahsulotni rasm bilan qayta ko'rsatamiz.
    if p:
        await _admin_send_product(message, p, pid)


# ─── Admin: mahsulot RASMLARINI o'zgartirish ─────────────────────────────────
def _admphoto_save_kb() -> InlineKeyboardMarkup:
    return InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="✅ Saqlash", callback_data="admphoto_save")],
        [InlineKeyboardButton(text="❌ Bekor qilish", callback_data="admphoto_cancel")],
    ])


@router.callback_query(F.data.startswith("admphoto_") & ~F.data.in_({"admphoto_save", "admphoto_cancel"}))
async def admin_edit_photo_start(call: CallbackQuery, state: FSMContext):
    if not is_owner(call.from_user.id): return
    pid = int(call.data.split("_")[1])
    await state.set_state(AdminEditProduct.waiting_photo)
    await state.update_data(pid=pid, new_photos=[])
    await call.message.answer(
        "🖼 Yangi rasm(lar)ni yuboring — albom qilib yoki bittalab.\n"
        "Tugagach «✅ Saqlash» tugmasini bosing.\n"
        "⚠️ Eski rasmlar yangilari bilan to'liq almashtiriladi."
    )
    await call.answer()


@router.message(AdminEditProduct.waiting_photo, F.media_group_id, F.photo)
async def admin_edit_photo_album(message: Message, state: FSMContext):
    key = (message.from_user.id, message.media_group_id)

    async def done(photos):
        d = await state.get_data()
        allp = (d.get("new_photos") or []) + photos
        await state.update_data(new_photos=allp)
        await message.answer(f"✅ {len(allp)} ta rasm qabul qilindi.", reply_markup=_admphoto_save_kb())

    collect(key, message.photo[-1].file_id, 1.5, done)


@router.message(AdminEditProduct.waiting_photo, F.photo)
async def admin_edit_photo_single(message: Message, state: FSMContext):
    data = await state.get_data()
    allp = (data.get("new_photos") or []) + [message.photo[-1].file_id]
    await state.update_data(new_photos=allp)
    await message.answer(f"✅ {len(allp)} ta rasm qabul qilindi.", reply_markup=_admphoto_save_kb())


@router.message(AdminEditProduct.waiting_photo)
async def admin_edit_photo_other(message: Message, state: FSMContext):
    if (message.text or "").strip() in ("/cancel", "❌ Bekor qilish"):
        await state.clear()
        await message.answer("Bekor qilindi.")
        return
    await message.answer("🖼 Iltimos, rasm yuboring (yoki «✅ Saqlash» / /cancel).")


@router.callback_query(AdminEditProduct.waiting_photo, F.data == "admphoto_save")
async def admin_edit_photo_save(call: CallbackQuery, state: FSMContext):
    data = await state.get_data()
    pid = data.get("pid")
    photos = data.get("new_photos") or []
    if not photos:
        await call.answer("Avval kamida 1 ta rasm yuboring.", show_alert=True); return
    update_product(pid, {"photos": photos})
    _log(call.from_user, "Mahsulot rasmlari o'zgartirildi", f"ID:{pid}, {len(photos)} ta rasm")
    await state.clear()
    await call.answer("✅ Rasmlar saqlandi")
    p = get_product_by_id(pid)
    if p:
        await _admin_send_product(call.message, p, pid)


@router.callback_query(AdminEditProduct.waiting_photo, F.data == "admphoto_cancel")
async def admin_edit_photo_cancel(call: CallbackQuery, state: FSMContext):
    data = await state.get_data()
    pid = data.get("pid")
    await state.clear()
    await call.answer("Bekor qilindi")
    p = get_product_by_id(pid)
    if p:
        await _admin_send_product(call.message, p, pid)


# ─── Admin: mahsulot VIDEOSINI o'zgartirish ──────────────────────────────────
@router.callback_query(F.data.startswith("admvideo_") & ~F.data.in_({"admvideo_del"}))
async def admin_edit_video_start(call: CallbackQuery, state: FSMContext):
    if not is_owner(call.from_user.id): return
    pid = int(call.data.split("_")[1])
    p = get_product_by_id(pid)
    await state.set_state(AdminEditProduct.waiting_video)
    await state.update_data(pid=pid)
    kb = None
    if p and p.get("video"):
        kb = InlineKeyboardMarkup(inline_keyboard=[
            [InlineKeyboardButton(text="🗑 Videoni o'chirish", callback_data="admvideo_del")],
        ])
    await call.message.answer(
        "🎬 Qisqa videoni yuboring (mahsulot kartochkasida ko'rinadi).\n"
        "Bekor qilish: /cancel",
        reply_markup=kb,
    )
    await call.answer()


@router.message(AdminEditProduct.waiting_video, F.video)
async def admin_edit_video_save(message: Message, state: FSMContext):
    if not is_owner(message.from_user.id):
        await state.clear(); return
    data = await state.get_data()
    pid = data.get("pid")
    update_product(pid, {"video": message.video.file_id})
    _log(message.from_user, "Mahsulot videosi o'zgartirildi", f"ID:{pid}")
    await state.clear()
    await message.answer("✅ Video saqlandi!")
    p = get_product_by_id(pid)
    if p:
        await _admin_send_product(message, p, pid)


@router.callback_query(AdminEditProduct.waiting_video, F.data == "admvideo_del")
async def admin_edit_video_del(call: CallbackQuery, state: FSMContext):
    data = await state.get_data()
    pid = data.get("pid")
    update_product(pid, {"video": None})
    _log(call.from_user, "Mahsulot videosi o'chirildi", f"ID:{pid}")
    await state.clear()
    await call.answer("🗑 Video o'chirildi")
    p = get_product_by_id(pid)
    if p:
        await _admin_send_product(call.message, p, pid)


@router.message(AdminEditProduct.waiting_video)
async def admin_edit_video_other(message: Message, state: FSMContext):
    if (message.text or "").strip() in ("/cancel", "❌ Bekor qilish"):
        await state.clear()
        await message.answer("Bekor qilindi.")
        return
    await message.answer("🎬 Iltimos, video yuboring (yoki /cancel).")


@router.callback_query(F.data.startswith("dprod_"))
async def admin_del_product(call: CallbackQuery):
    if not is_owner(call.from_user.id): return
    pid = int(call.data.split("_")[1])
    _p = get_product_by_id(pid)
    admin_delete_product(pid)
    _log(call.from_user, "Mahsulot o'chirildi",
         f"{_p['name']} (ID:{pid})" if _p else f"ID:{pid}")
    name = _p["name"] if _p else f"#{pid}"
    if get_all_products():
        page_text, kb = _products_page(0)
        text = f"🗑 <b>{name}</b> o'chirildi.\n\n" + page_text
    else:
        kb = InlineKeyboardMarkup(inline_keyboard=[[InlineKeyboardButton(text="🔙 Orqaga", callback_data="admin_back")]])
        text = f"🗑 <b>{name}</b> o'chirildi.\n\nMahsulot yo'q."
    await _admin_nav(call, text, kb)
    await call.answer("O'chirildi")



# ─── Foydalanuvchilar ────────────────────────────────────────────────────────
@router.callback_query(F.data == "admin_users")
async def admin_users(call: CallbackQuery):
    if not is_owner(call.from_user.id): return
    await _ack(call)
    from app.storage import get_users
    users = get_users()
    sellers = get_sellers()
    text = f"👥 <b>Foydalanuvchilar:</b> {len(users)} ta\n🏪 Sellerlar: {len(sellers)} ta\n\n"
    rows = []
    for uid, u in list(users.items())[:20]:
        role = "🏪" if uid in sellers else "🛍"
        rows.append([InlineKeyboardButton(
            text=f"{role} {u.get('full_name','ID:'+uid)}",
            callback_data=f"auser_{uid}"
        )])
    rows.append([InlineKeyboardButton(text="🔙 Orqaga", callback_data="admin_back")])
    await call.message.edit_text(text, parse_mode="HTML",
                                  reply_markup=InlineKeyboardMarkup(inline_keyboard=rows))


@router.callback_query(F.data.startswith("auser_"))
async def admin_user_detail(call: CallbackQuery):
    if not is_owner(call.from_user.id): return
    uid = call.data.split("_")[1]
    from app.storage import get_users
    u = get_users().get(uid, {})
    sellers = get_sellers()
    role = "🏪 Seller" if uid in sellers else "🛍 Xaridor"
    text = f"👤 <b>{u.get('full_name','Noma lum')}</b>\nID: {uid}\nRol: {role}"
    kb = InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="🗑 O'chirish", callback_data=f"del_user_{uid}")],
        [InlineKeyboardButton(text="🔙 Orqaga",    callback_data="admin_users")],
    ])
    await call.message.edit_text(text, parse_mode="HTML", reply_markup=kb)
    await call.answer()


@router.callback_query(F.data.startswith("del_user_"))
async def del_user_cb(call: CallbackQuery):
    if not is_owner(call.from_user.id): return
    uid = int(call.data.split("_")[2])
    delete_user(uid)
    delete_seller(uid)
    _log(call.from_user, "Foydalanuvchi o'chirildi", f"ID:{uid}")
    await call.message.edit_text(f"🗑 Foydalanuvchi (ID:{uid}) o'chirildi.")
    await call.answer("O'chirildi")


# ─── Botni bloklaganlar ──────────────────────────────────────────────────────
# Telegram foydalanuvchi botni bloklaganda/blokdan chiqarganda "my_chat_member"
# yangilanishini yuboradi. Shu yerda ushlab, bloklaganlar ro'yxatini yuritamiz.
@router.my_chat_member()
async def on_my_chat_member(event: ChatMemberUpdated):
    # Faqat shaxsiy chatdagi (foydalanuvchi ↔ bot) holat o'zgarishi muhim.
    if event.chat.type != "private":
        return
    u = event.from_user
    status = event.new_chat_member.status
    if status == "kicked":
        # Foydalanuvchi botni BLOKLADI
        mark_blocked(u.id, {
            "user_id":  u.id,
            "name":     u.full_name,
            "username": u.username or "",
        })
        # Owner'ga darhol xabar beramiz
        try:
            from app.bot.bot import bot
            uname = f"@{u.username}" if u.username else "—"
            await bot.send_message(
                settings.OWNER_ID,
                f"🚫 <b>Foydalanuvchi botni bloklab qo'ydi</b>\n\n"
                f"👤 {u.full_name}\n"
                f"🔗 Username: {uname}\n"
                f"🆔 ID: <code>{u.id}</code>",
                parse_mode="HTML",
            )
        except Exception:
            pass
    elif status == "member":
        # Foydalanuvchi botni blokdan chiqardi — ro'yxatdan olib tashlaymiz
        unmark_blocked(u.id)


@router.callback_query(F.data == "admin_blocked")
async def admin_blocked(call: CallbackQuery):
    if not is_owner(call.from_user.id): return
    await _ack(call)
    blocked = get_blocked()
    if not blocked:
        await _admin_nav(
            call,
            "🚫 <b>Bloklaganlar</b>\n\nHozircha botni bloklagan foydalanuvchi yo'q. ✅",
            InlineKeyboardMarkup(inline_keyboard=[
                [InlineKeyboardButton(text="🔙 Orqaga", callback_data="admin_back")]
            ]),
        )
        return
    # Oxirgi bloklaganlar yuqorida tursin
    items = sorted(blocked.values(), key=lambda b: b.get("blocked_at", ""), reverse=True)
    lines = [f"🚫 <b>Botni bloklaganlar</b> — jami {len(items)} ta\n"]
    for i, b in enumerate(items, 1):
        uname = f"@{b['username']}" if b.get("username") else "—"
        when = (b.get("blocked_at", "") or "")[:16].replace("T", " ")
        lines.append(
            f"<b>{i}. {b.get('name','—')}</b>\n"
            f"   🔗 {uname}\n"
            f"   🆔 <code>{b.get('user_id','—')}</code>\n"
            f"   🕒 {when}"
        )
    await _admin_nav(
        call,
        "\n".join(lines),
        InlineKeyboardMarkup(inline_keyboard=[
            [InlineKeyboardButton(text="🔙 Orqaga", callback_data="admin_back")]
        ]),
    )


# ─── Statistika ──────────────────────────────────────────────────────────────
@router.callback_query(F.data == "admin_stats")
async def show_stats(call: CallbackQuery):
    if not is_owner(call.from_user.id): return
    await _ack(call)
    apps = get_applications()
    orders = get_orders()
    revenue = sum(o.get("total", 0) for o in orders)
    text = (
        f"📊 <b>Statistika</b>\n\n"
        f"📋 Arizalar: {len(apps)} ta\n"
        f"✅ Tasdiqlangan: {len([a for a in apps.values() if a.get('status')=='approved'])}\n"
        f"🏪 Sellerlar: {len(get_sellers())}\n"
        f"📦 Mahsulotlar: {len(get_all_products())}\n"
        f"🛒 Buyurtmalar: {len(orders)} ta\n"
        f"💰 Jami savdo: {revenue:,} so'm"
    )
    await call.message.edit_text(text, parse_mode="HTML", reply_markup=InlineKeyboardMarkup(
        inline_keyboard=[[InlineKeyboardButton(text="🔙 Orqaga", callback_data="admin_back")]]
    ))


# ─── Excel hisobot ───────────────────────────────────────────────────────────
@router.callback_query(F.data == "admin_excel_menu")
async def excel_menu(call: CallbackQuery):
    if not is_owner(call.from_user.id): return
    kb = InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="📅 Kunlik hisobot",  callback_data="excel_daily")],
        [InlineKeyboardButton(text="📆 Oylik hisobot",   callback_data="excel_monthly")],
        [InlineKeyboardButton(text="🏪 Seller bo'yicha", callback_data="excel_sellers")],
        [InlineKeyboardButton(text="🔙 Orqaga",          callback_data="admin_back")],
    ])
    await call.message.edit_text("📈 Qaysi hisobotni yuklab olmoqchisiz?", reply_markup=kb)
    await call.answer()


def _build_excel(orders: list, title: str) -> bytes:
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = title[:30]

    headers = ["#", "Sana", "Mahsulot", "Xaridor", "Telefon", "Manzil",
               "Yetkazish", "Narx", "Komissiya 10%", "Holat"]
    bold = Font(bold=True, color="FFFFFF")
    fill = PatternFill("solid", fgColor="2E86AB")
    for col, h in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=h)
        cell.font = bold
        cell.fill = fill
        cell.alignment = Alignment(horizontal="center")

    dlv_map = {"pickup": "O'zi olib ketadi", "taxi": "Taksi pochta (shu bugunoq)", "btc": "BTC Pochta", "emu": "EMU Express", "uzum": "Uzum Pochta"}
    total = 0
    total_comm = 0
    for i, o in enumerate(orders, 1):
        comm = o.get("commission", int(o.get("total", 0) * 0.1))
        ws.append([
            o.get("id", i),
            o.get("created_at", "")[:10],
            o.get("product_name", "—"),
            o.get("buyer_name", o.get("buyer_id", "—")),
            o.get("phone", "—"),
            o.get("address", "—"),
            dlv_map.get(o.get("delivery", ""), o.get("delivery", "—")),
            o.get("total", 0),
            comm,
            o.get("status", "—"),
        ])
        total += o.get("total", 0)
        total_comm += comm

    ws.append([])
    ws.append(["", "", "", "", "", "", "JAMI:", total, total_comm, ""])
    for c in (7, 8, 9):
        ws.cell(row=ws.max_row, column=c).font = Font(bold=True)

    widths = [6, 12, 22, 20, 16, 30, 14, 14, 14, 14]
    for idx, w in enumerate(widths, 1):
        ws.column_dimensions[ws.cell(row=1, column=idx).column_letter].width = w

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


@router.callback_query(F.data == "excel_daily")
async def excel_daily(call: CallbackQuery):
    if not is_owner(call.from_user.id): return
    await _ack(call)
    today = datetime.now().date().isoformat()
    orders = [o for o in get_orders() if o.get("created_at", "")[:10] == today]
    data = _build_excel(orders, "Kunlik")
    file = BufferedInputFile(data, filename=f"kunlik_{today}.xlsx")
    await call.message.answer_document(file, caption=f"📅 Kunlik hisobot — {today}\nBuyurtmalar: {len(orders)} ta")


@router.callback_query(F.data == "excel_monthly")
async def excel_monthly(call: CallbackQuery):
    if not is_owner(call.from_user.id): return
    await _ack(call)
    month = datetime.now().strftime("%Y-%m")
    orders = [o for o in get_orders() if o.get("created_at", "")[:7] == month]
    data = _build_excel(orders, "Oylik")
    file = BufferedInputFile(data, filename=f"oylik_{month}.xlsx")
    await call.message.answer_document(file, caption=f"📆 Oylik hisobot — {month}\nBuyurtmalar: {len(orders)} ta")


# ─── Seller bo'yicha hisobot ─────────────────────────────────────────────────
@router.callback_query(F.data == "excel_sellers")
async def excel_sellers_list(call: CallbackQuery):
    if not is_owner(call.from_user.id): return
    sellers = get_sellers()
    if not sellers:
        await call.answer("Seller yo'q.", show_alert=True); return
    await _ack(call)
    # Buyurtma sonini bitta o'tishda yig'amiz (har seller uchun alohida emas).
    cnt_by_seller: dict = {}
    for o in get_orders():
        sid = o.get("seller_id")
        cnt_by_seller[sid] = cnt_by_seller.get(sid, 0) + 1
    rows = []
    for uid, s in sellers.items():
        cnt = cnt_by_seller.get(int(uid), 0)
        rows.append([InlineKeyboardButton(
            text=f"🏪 {s['shop_name']} ({cnt} buyurtma)",
            callback_data=f"excel_seller_{uid}"
        )])
    rows.append([InlineKeyboardButton(text="🔙 Orqaga", callback_data="admin_excel_menu")])
    await call.message.edit_text(
        "🏪 <b>Qaysi seller hisoboti?</b>", parse_mode="HTML",
        reply_markup=InlineKeyboardMarkup(inline_keyboard=rows)
    )


@router.callback_query(F.data.startswith("excel_seller_"))
async def excel_seller_report(call: CallbackQuery):
    if not is_owner(call.from_user.id): return
    uid = int(call.data.split("_")[-1])
    seller = get_seller(uid)
    orders = get_seller_orders(uid)
    if not orders:
        await call.answer("Bu sellerda buyurtma yo'q.", show_alert=True); return
    await _ack(call)
    shop = seller["shop_name"] if seller else str(uid)
    data = _build_excel(orders, f"{shop}")
    total = sum(o.get("total", 0) for o in orders)
    comm  = sum(o.get("commission", int(o.get("total", 0) * 0.1)) for o in orders)
    file = BufferedInputFile(data, filename=f"seller_{uid}.xlsx")
    await call.message.answer_document(
        file,
        caption=(
            f"🏪 <b>{shop}</b> — hisobot\n"
            f"🛒 Buyurtmalar: {len(orders)} ta\n"
            f"💰 Jami savdo: {total:,} so'm\n"
            f"💵 Sizning komissiyangiz (10%): {comm:,} so'm"
        ),
        parse_mode="HTML"
    )


# ═══════════════════════════════════════════════════════════════════════════
#  ADMIN: MAHSULOT QO'SHISH (owner + sub-admin)
# ═══════════════════════════════════════════════════════════════════════════
class AdminAddState(StatesGroup):
    user_id = State()
    city    = State()


@router.callback_query(F.data == "admin_addprod")
async def admin_addprod_pick(call: CallbackQuery):
    if not is_admin(call.from_user.id): return
    sellers = list(get_sellers().items())
    if not is_owner(call.from_user.id):
        mycity = _admin_city(call.from_user.id)
        sellers = [(u, s) for u, s in sellers if s.get("city") == mycity]
    if not sellers:
        await call.answer("Sizning shahringizda do'kon yo'q.", show_alert=True); return
    rows = [[InlineKeyboardButton(text=f"🏪 {s['shop_name']} ({s.get('city','—')})",
                                  callback_data=f"aap_{uid}")]
            for uid, s in sellers]
    rows.append([InlineKeyboardButton(text="🔙 Orqaga", callback_data="admin_back")])
    await _admin_nav(call, "➕ <b>Mahsulot qaysi do'konga qo'shilsin?</b>",
                     InlineKeyboardMarkup(inline_keyboard=rows))
    await call.answer()


@router.callback_query(F.data.startswith("aap_"))
async def admin_addprod_start(call: CallbackQuery, state: FSMContext):
    if not is_admin(call.from_user.id): return
    sid = int(call.data.split("_")[1])
    seller = get_seller(sid)
    if not seller:
        await call.answer("Do'kon topilmadi.", show_alert=True); return
    await state.set_state(AdminAddProduct.name)
    await state.update_data(seller_id=sid, shop_name=seller["shop_name"], city=seller.get("city", ""))
    await call.message.answer(f"🏪 {seller['shop_name']}\n📦 Mahsulot nomini kiriting:")
    await call.answer()


# ─── Mahsulot qo'shishni bo'lib yuboruvchi tugma/buyruqlarda AVTOMAT to'xtatish ──
_ADMIN_ADD_STATES = StateFilter(
    AdminAddProduct.name, AdminAddProduct.description,
    AdminAddProduct.price, AdminAddProduct.photo,
)


@router.message(_ADMIN_ADD_STATES, F.text.startswith("/"), F.text != "/skip")
async def admin_ap_interrupt_command(message: Message, state: FSMContext):
    if (message.text or "").startswith("/start"):
        from app.handlers.start import cmd_start
        await cmd_start(message, state)
        return
    await state.clear()
    await message.answer("⛔️ Mahsulot qo'shish to'xtatildi.", reply_markup=main_menu)


@router.message(_ADMIN_ADD_STATES, F.text.in_(_MENU_BUTTONS))
async def admin_ap_interrupt_menu(message: Message, state: FSMContext):
    await state.clear()
    await message.answer("⛔️ Mahsulot qo'shish to'xtatildi.", reply_markup=main_menu)


@router.message(AdminAddProduct.name)
async def admin_ap_name(message: Message, state: FSMContext):
    if not (message.text or "").strip():
        await message.answer("❌ Mahsulot nomini matn ko'rinishida kiriting:"); return
    await state.update_data(name=message.text.strip())
    await state.set_state(AdminAddProduct.description)
    await message.answer("📝 Tavsif kiriting:")


@router.message(AdminAddProduct.description)
async def admin_ap_desc(message: Message, state: FSMContext):
    if not (message.text or "").strip():
        await message.answer("❌ Tavsifni matn ko'rinishida kiriting:"); return
    await state.update_data(description=message.text.strip())
    await state.set_state(AdminAddProduct.price)
    await message.answer("💰 Narxini kiriting (faqat raqam):")


@router.message(AdminAddProduct.price)
async def admin_ap_price(message: Message, state: FSMContext):
    if not (message.text or "").isdigit():
        await message.answer("❌ Faqat raqam kiriting:"); return
    await state.update_data(price=int(message.text))
    await state.set_state(AdminAddProduct.photo)
    await message.answer(
        "📸 Rasm(lar)ni yuboring — bittasini yoki bir nechtasini birga "
        "(albom) jo'nating.\nRasmsiz qo'shish uchun /skip yozing."
    )


def _admin_build_product(data: dict, photos: list) -> dict:
    return {
        "seller_id":   data["seller_id"],
        "shop_name":   data["shop_name"],
        "city":        data.get("city", ""),
        "name":        data["name"],
        "description": data["description"],
        "price":       data["price"],
        "photos":      photos,
    }


async def _admin_finish_product(message: Message, state: FSMContext, data: dict, photos: list):
    add_product(_admin_build_product(data, photos))
    _log(message.from_user, "Mahsulot qo'shildi",
         f"{data['name']} — do'kon: {data['shop_name']} ({data['price']:,} so'm)")
    await state.clear()
    await message.answer(
        f"✅ <b>{data['name']}</b> «{data['shop_name']}» do'koniga qo'shildi!"
        + (f"\n🖼 {len(photos)} ta rasm" if photos else ""),
        parse_mode="HTML", reply_markup=admin_menu_kb(message.from_user.id)
    )
    try:
        from app.bot.bot import bot
        await bot.send_message(
            data["seller_id"],
            f"➕ Admin do'koningizga yangi mahsulot qo'shdi:\n"
            f"📦 {data['name']} — {data['price']:,} so'm",
        )
    except Exception:
        pass


@router.message(AdminAddProduct.photo, F.media_group_id)
async def admin_ap_photo_album(message: Message, state: FSMContext):
    data = await state.get_data()
    key  = (message.from_user.id, message.media_group_id)

    async def done(photos):
        await _admin_finish_product(message, state, data, photos)

    collect(key, message.photo[-1].file_id, 1.5, done)


@router.message(AdminAddProduct.photo, F.photo)
async def admin_ap_photo_single(message: Message, state: FSMContext):
    data = await state.get_data()
    await _admin_finish_product(message, state, data, [message.photo[-1].file_id])


@router.message(AdminAddProduct.photo, F.text == "/skip")
async def admin_ap_photo_skip(message: Message, state: FSMContext):
    data = await state.get_data()
    await _admin_finish_product(message, state, data, [])


@router.message(AdminAddProduct.photo)
async def admin_ap_photo_wrong(message: Message, state: FSMContext):
    await message.answer("❌ Rasm yuboring yoki /skip yozing:")


# ═══════════════════════════════════════════════════════════════════════════
#  OWNER: SUB-ADMINLARNI BOSHQARISH
# ═══════════════════════════════════════════════════════════════════════════
@router.callback_query(F.data == "admin_admins")
async def admins_list(call: CallbackQuery):
    if not is_owner(call.from_user.id): return
    admins = get_admins()
    text = f"👮 <b>Sub-adminlar</b> — {len(admins)} ta\n\n"
    rows = []
    if not admins:
        text += "Hozircha sub-admin yo'q.\n"
    for uid, a in admins.items():
        text += f"• {a.get('name','—')} — 🏙 {a.get('city','—')} (ID: {uid})\n"
        rows.append([InlineKeyboardButton(text=f"🗑 {a.get('name', uid)} ni olib tashlash",
                                          callback_data=f"deladmin_{uid}")])
    rows.append([InlineKeyboardButton(text="➕ Admin qo'shish", callback_data="addadmin")])
    rows.append([InlineKeyboardButton(text="🔙 Orqaga", callback_data="admin_back")])
    text += ("\n<i>Sub-admin faqat seller va mahsulot qo'sha oladi. "
             "Uning har bir amali jurnalga yoziladi.</i>")
    await _admin_nav(call, text, InlineKeyboardMarkup(inline_keyboard=rows))
    await call.answer()


@router.callback_query(F.data == "addadmin")
async def add_admin_start(call: CallbackQuery, state: FSMContext):
    if not is_owner(call.from_user.id): return
    await state.set_state(AdminAddState.user_id)
    await call.message.answer(
        "➕ Yangi adminning <b>Telegram ID</b> raqamini yuboring.\n"
        "(Foydalanuvchi avval botga /start bosgan bo'lishi kerak. "
        "ID'ni u 👤 Profil bo'limida ko'radi.)",
        parse_mode="HTML"
    )
    await call.answer()


@router.message(AdminAddState.user_id)
async def add_admin_save(message: Message, state: FSMContext):
    if not is_owner(message.from_user.id):
        await state.clear(); return
    txt = (message.text or "").strip()
    if not txt.isdigit():
        await message.answer("❌ Faqat raqamli ID yuboring:"); return
    uid = int(txt)
    if uid == settings.OWNER_ID:
        await state.clear()
        await message.answer("Bu siz — asosiy egasiz."); return
    u = get_user(uid)
    name = u.get("full_name") if u else f"ID {uid}"
    await state.update_data(new_admin_id=uid, new_admin_name=name)
    await state.set_state(AdminAddState.city)
    rows, row = [], []
    for c in get_cities():
        row.append(InlineKeyboardButton(text=c, callback_data=f"admincity_{c}"))
        if len(row) == 2:
            rows.append(row); row = []
    if row:
        rows.append(row)
    await message.answer(
        f"👤 {name}\n🏙 Bu admin qaysi shaharni boshqaradi? Tanlang:",
        reply_markup=InlineKeyboardMarkup(inline_keyboard=rows)
    )


@router.callback_query(AdminAddState.city, F.data.startswith("admincity_"))
async def add_admin_city(call: CallbackQuery, state: FSMContext):
    if not is_owner(call.from_user.id): return
    city = call.data.split("_", 1)[1]
    if city not in get_cities():
        await call.answer("Shahar topilmadi.", show_alert=True); return
    data = await state.get_data()
    uid = data["new_admin_id"]; name = data["new_admin_name"]
    await state.clear()
    add_admin(uid, {"name": name, "added_by": call.from_user.id, "city": city})
    _log(call.from_user, "Sub-admin qo'shildi", f"{name} (ID:{uid}, shahar: {city})")
    await call.message.answer(
        f"✅ {name} — <b>{city}</b> shahri admini qilib tayinlandi.\n"
        f"U /admin orqali faqat shu shahar seller/mahsulotlarini boshqaradi.",
        parse_mode="HTML", reply_markup=admin_menu_kb(call.from_user.id)
    )
    try:
        from app.bot.bot import bot
        await bot.send_message(uid, f"🛡 Sizga <b>{city}</b> shahri admini huquqi berildi!\n/admin buyrug'ini bosing.",
                               parse_mode="HTML")
    except Exception:
        pass
    await call.answer("Qo'shildi!")


@router.callback_query(F.data.startswith("deladmin_"))
async def del_admin_cb(call: CallbackQuery):
    if not is_owner(call.from_user.id): return
    uid = int(call.data.split("_")[1])
    a = get_admins().get(str(uid), {})
    if remove_admin(uid):
        _log(call.from_user, "Sub-admin olib tashlandi", f"{a.get('name','—')} (ID:{uid})")
        await call.answer("Olib tashlandi.")
    else:
        await call.answer("Topilmadi.")
    await admins_list(call)


# ═══════════════════════════════════════════════════════════════════════════
#  OWNER: KURIERLAR BOSHQARUVI
#  To'lov (10% oldindan) tasdiqlangan yetkazib berish zakazlari shu
#  kurierlarga avtomatik yuboriladi (confirm_payment ichida).
# ═══════════════════════════════════════════════════════════════════════════
class AdminCourierState(StatesGroup):
    user_id = State()
    name    = State()


@router.callback_query(F.data == "admin_couriers")
async def couriers_list(call: CallbackQuery):
    if not is_owner(call.from_user.id): return
    await _ack(call)
    couriers = get_couriers()
    text = f"🚚 <b>Kurierlar</b> — jami {len(couriers)} ta\n\n"
    rows = []
    if not couriers:
        text += "Hozircha kurier yo'q.\n"
    for uid, c in couriers.items():
        text += f"• {c.get('name','—')} (ID: {uid})\n"
        rows.append([InlineKeyboardButton(text=f"🗑 {c.get('name', uid)} ni olib tashlash",
                                          callback_data=f"delcourier_{uid}")])
    rows.append([InlineKeyboardButton(text="➕ Kurier qo'shish", callback_data="addcourier")])
    rows.append([InlineKeyboardButton(text="🔙 Orqaga", callback_data="admin_back")])
    text += ("\n<i>To'lov (10%) tasdiqlangan har bir yetkazib berish zakazi "
             "shu kurierlarga avtomatik yuboriladi. Kurier botga /start "
             "bosgan bo'lishi shart.</i>")
    await _admin_nav(call, text, InlineKeyboardMarkup(inline_keyboard=rows))


@router.callback_query(F.data == "addcourier")
async def add_courier_start(call: CallbackQuery, state: FSMContext):
    if not is_owner(call.from_user.id): return
    await _ack(call)
    await state.set_state(AdminCourierState.user_id)
    await call.message.answer(
        "➕ Kurierning <b>Telegram ID</b> raqamini yuboring.\n"
        "(Kurier avval botga /start bosgan bo'lishi kerak. "
        "ID'ni u 👤 Profil bo'limida ko'radi.)",
        parse_mode="HTML"
    )


@router.message(AdminCourierState.user_id)
async def add_courier_id(message: Message, state: FSMContext):
    if not is_owner(message.from_user.id):
        await state.clear(); return
    txt = (message.text or "").strip()
    if not txt.isdigit():
        await message.answer("❌ Faqat raqamli ID yuboring:"); return
    uid = int(txt)
    u = get_user(uid)
    name = u.get("full_name") if u else ""
    await state.update_data(courier_id=uid, courier_name=name)
    await state.set_state(AdminCourierState.name)
    hint = f" (botda: {name})" if name else ""
    await message.answer(f"👤 Kurierning ismini yozing{hint}:")


@router.message(AdminCourierState.name)
async def add_courier_name(message: Message, state: FSMContext):
    if not is_owner(message.from_user.id):
        await state.clear(); return
    data = await state.get_data()
    await state.clear()
    uid  = data["courier_id"]
    name = (message.text or "").strip() or data.get("courier_name") or "Kurier"
    add_courier(uid, {"user_id": uid, "name": name})
    _log(message.from_user, "Kurier qo'shildi", f"{name} (ID:{uid})")
    try:
        from app.bot.bot import bot
        await bot.send_message(
            uid,
            "🚚 Siz kurier sifatida qo'shildingiz!\n"
            "To'lovi tasdiqlangan yangi zakazlar sizga shu yerga keladi.",
        )
    except Exception:
        pass
    await message.answer(
        f"✅ Kurier qo'shildi: <b>{name}</b> (ID: {uid})",
        parse_mode="HTML", reply_markup=admin_menu_kb(message.from_user.id)
    )


@router.callback_query(F.data.startswith("delcourier_"))
async def del_courier_cb(call: CallbackQuery):
    if not is_owner(call.from_user.id): return
    uid = call.data.split("_")[1]
    c = get_couriers().get(uid, {})
    if remove_courier(uid):
        _log(call.from_user, "Kurier olib tashlandi", f"{c.get('name','—')} (ID:{uid})")
        await call.answer("Olib tashlandi.")
    else:
        await call.answer("Topilmadi.")
    await couriers_list(call)


def _courier_order_text(oid: int, o: dict) -> str:
    """Kurierga yuboriladigan zakaz kartochkasi (yangi zakaz xabari va /kurier
    panelida bir xil ishlatiladi)."""
    shop = get_seller(o["seller_id"]) or {}
    total  = o.get("total", 0)
    prepay = o.get("prepay", 0)
    remain = max(total - prepay, 0)
    fee    = o.get("delivery_fee", 0)
    fee_line = (
        f"🚚 Yetkazish haqi: <b>{fee:,} so'm</b> (xaridordan olasiz)\n" if fee else ""
    )
    return (
        f"🚚 <b>ZAKAZ</b>\n\n"
        f"🆔 Zakaz raqami: <b>#{oid}</b>\n"
        f"🏪 Do'kon: <b>{shop.get('shop_name','—')}</b>\n"
        f"📱 Do'kon tel: {shop.get('phone','—')}\n"
        f"🏙 Shahar: {shop.get('city','—')}\n"
        f"📦 Mahsulot: {o.get('product_name','—')}\n"
        f"💰 Qoldiq summa: <b>{remain:,} so'm</b>"
        f"  (jami: {total:,}, oldindan to'langan: {prepay:,})\n"
        f"{fee_line}\n"
        f"📍 Yetkazish manzili: {o.get('address','—')}\n"
        f"👤 Xaridor: {o.get('buyer_name','—')} — {o.get('phone','—')}\n\n"
        f"Do'kondan mahsulotni olib, xaridorga yetkazing."
    )


def _courier_done_kb(oid: int) -> InlineKeyboardMarkup:
    return InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="✅ Yetkazib berildi", callback_data=f"cdone_{oid}")]
    ])


# ─── Kurier paneli: faol zakazlar ro'yxati ───────────────────────────────────
@router.message(Command("kurier"))
async def courier_panel(message: Message):
    if not (is_courier(message.from_user.id) or is_owner(message.from_user.id)):
        await message.answer("❌ Siz kurier emassiz.")
        return
    active = [
        o for o in get_orders()
        if o.get("delivery") != "pickup"
        and o.get("status") in ("paid", "processing", "shipped")
    ]
    if not active:
        await message.answer("🚚 Hozircha faol zakaz yo'q.")
        return
    await message.answer(f"🚚 <b>Faol zakazlar:</b> {len(active)} ta", parse_mode="HTML")
    for o in active[-10:]:
        await message.answer(_courier_order_text(o["id"], o), parse_mode="HTML",
                             reply_markup=_courier_done_kb(o["id"]))


# ─── Kurier "Yetkazib berildi" → xaridor qolgan 90% + yetkazishni KURIERGA naqd beradi ──
# Yangi oqim: kurier manzilga yetib tugmani bosadi → xaridorga "qolgan to'lov +
# yetkazish haqini kurierга naqd bering" xabari boradi. Karta kerak bo'lsa xaridor
# kurierdan so'raydi (kurierда seller kartasi bor). Pulni olgach kurier
# "Pulni oldim — topshirdim" tugmasini bosadi va zakaz yakunlanadi.


def _courier_confirm_kb(oid: int) -> InlineKeyboardMarkup:
    return InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="🤝 Pulni oldim — topshirdim",
                              callback_data=f"cpaid_{oid}")]
    ])


async def _finalize_handover(oid: int, o: dict):
    """Zakaz yakunlanadi: holat 'delivered', xaridorga baholash, sellerga xabar.
    Qolgan to'lovni kurier naqd oldi (yoki qoldiq yo'q edi)."""
    from app.bot.bot import bot
    update_order_status(oid, "delivered")
    total  = o.get("total", 0)
    prepay = o.get("prepay", 0)
    remain = max(total - prepay, 0)

    # Xaridorga — yakun + baholash
    try:
        await bot.send_message(
            o["buyer_id"],
            f"📦 <b>Buyurtmangiz topshirildi!</b>  (#{oid})\n"
            f"📦 {o.get('product_name', '—')}\n\n"
            "⭐ Sellerni baholang:",
            parse_mode="HTML",
            reply_markup=stars_kb(o["seller_id"], oid)
        )
    except Exception:
        pass

    # ── Sellerga (ega + yordamchilar) — MAHSULOT TOPSHIRILDI ──
    seller_done = (
        f"✅ <b>Buyurtma #{oid} — mahsulot xaridorga topshirildi!</b>\n"
        f"📦 {o.get('product_name','—')}\n"
    )
    if remain > 0:
        seller_done += (
            f"💰 Qolgan to'lov ({remain:,} so'm) kurierда — kurier siz bilan "
            f"hisob-kitob qiladi.\n"
        )
    for nid in shop_notify_ids(o["seller_id"]):
        try:
            await bot.send_message(nid, seller_done, parse_mode="HTML")
        except Exception:
            pass


@router.callback_query(F.data.startswith("cdone_"))
async def courier_delivered(call: CallbackQuery):
    if not (is_courier(call.from_user.id) or is_owner(call.from_user.id)):
        await call.answer("Ruxsat yo'q.", show_alert=True); return
    oid = int(call.data.split("_")[1])
    o = get_order_by_id(oid)
    if not o:
        await call.answer("Buyurtma topilmadi.", show_alert=True); return
    if o.get("status") == "delivered":
        await call.answer("Bu zakaz allaqachon yetkazilgan.", show_alert=True); return

    total  = o.get("total", 0)
    prepay = o.get("prepay", 0)
    remain = max(total - prepay, 0)
    fee    = o.get("delivery_fee", 0)
    from app.bot.bot import bot
    courier = get_couriers().get(str(call.from_user.id), {})

    # Qolgan to'lov bo'lmasa — tasdiqlashsiz darhol yakunlanadi
    if remain <= 0:
        await call.answer("✅ Yetkazib berildi!")
        _log(call.from_user, "Zakaz yetkazildi (kurier)",
             f"Buyurtma #{oid} — kurier: {courier.get('name', call.from_user.full_name)}")
        try:
            await call.message.edit_text(
                call.message.html_text + "\n\n✅ <b>YETKAZIB BERILDI</b>",
                parse_mode="HTML",
            )
        except Exception:
            pass
        await _finalize_handover(oid, o)
        return

    # Takror bosilgan bo'lsa — faqat tasdiqlash tugmasini qayta ko'rsatamiz
    if o.get("handover_at"):
        await call.answer("⏳ Pulni olib, tasdiqlang.", show_alert=True)
        try:
            await call.message.edit_reply_markup(reply_markup=_courier_confirm_kb(oid))
        except Exception:
            pass
        return

    update_order_fields(oid, {
        "handover_at": datetime.now().isoformat(),
        "handover_courier_id": call.from_user.id,
    })
    await call.answer("📍 Qabul qilindi — pulni oling va tasdiqlang.")
    _log(call.from_user, "Zakaz manzilga yetkazildi (kurier)",
         f"Buyurtma #{oid} — kurier: {courier.get('name', call.from_user.full_name)}")

    due = remain + fee
    fee_txt = f"{fee:,} so'm" if fee else "bepul"
    # Xaridorga — qolgan to'lov + yetkazish KURIERGA naqd beriladi
    try:
        await bot.send_message(
            o["buyer_id"],
            f"<b>Buyurtmangiz yetib keldi</b>  (#{oid})\n"
            f"{o.get('product_name','—')}\n\n"
            f"Qolgan to'lov:  <b>{remain:,} so'm</b>\n"
            f"Yetkazib berish:  <b>{fee_txt}</b>\n"
            f"Jami kurierga (naqd):  <b>{due:,} so'm</b>\n\n"
            "Mahsulotni tekshirib, summani <b>kurierga naqd</b> bering.\n"
            "Karta orqali to'lamoqchi bo'lsangiz — kurierdan karta raqamini so'rang.",
            parse_mode="HTML"
        )
    except Exception:
        pass

    # Kurierга — naqd oladigan summa + (so'rashsa beradigan) seller kartasi + tugma
    seller_card = (get_seller(o["seller_id"]) or {}).get("card_number", "")
    card_hint = (
        f"\n💳 Karta so'rashsa bering: <code>{seller_card}</code>" if seller_card else ""
    )
    try:
        await call.message.edit_text(
            call.message.html_text +
            f"\n\n💵 <b>Xaridordan naqd oling: {due:,} so'm</b>"
            f"  (mahsulot {remain:,}{f' + yetkazish {fee:,}' if fee else ', yetkazish bepul'})"
            f"{card_hint}"
            f"\nPulni olib, mahsulotni topshirgach quyidagi tugmani bosing 👇",
            parse_mode="HTML",
            reply_markup=_courier_confirm_kb(oid)
        )
    except Exception:
        pass


# ─── Kurier: "Pulni oldim — topshirdim" ──────────────────────────────────────
@router.callback_query(F.data.startswith("cpaid_"))
async def courier_payment_confirm(call: CallbackQuery):
    if not (is_courier(call.from_user.id) or is_owner(call.from_user.id)):
        await call.answer("Ruxsat yo'q.", show_alert=True); return
    oid = int(call.data.split("_")[1])
    o = get_order_by_id(oid)
    if not o:
        await call.answer("Buyurtma topilmadi.", show_alert=True); return
    if o.get("status") == "delivered":
        await call.answer("Allaqachon tasdiqlangan.", show_alert=True); return

    await call.answer("🤝 Tasdiqlandi — rahmat!")
    courier = get_couriers().get(str(call.from_user.id), {})
    _log(call.from_user, "Qolgan to'lov kurier tomonidan olindi",
         f"Buyurtma #{oid} — kurier: {courier.get('name', call.from_user.full_name)}")
    try:
        await call.message.edit_text(
            call.message.html_text + "\n\n🤝 <b>PUL OLINDI — TOPSHIRILDI</b>",
            parse_mode="HTML",
        )
    except Exception:
        pass
    await _finalize_handover(oid, o)


# ═══════════════════════════════════════════════════════════════════════════
#  OWNER: AUDIT JURNALI + WORD
# ═══════════════════════════════════════════════════════════════════════════
def _fmt_dt(iso: str) -> str:
    try:
        return datetime.fromisoformat(iso).strftime("%d.%m.%Y %H:%M")
    except Exception:
        return iso[:16]


@router.callback_query(F.data == "admin_log")
async def admin_log(call: CallbackQuery):
    if not is_owner(call.from_user.id): return
    log = get_audit()
    text = f"📜 <b>Amallar jurnali</b> — jami {len(log)} ta\n\n"
    if not log:
        text += "Hozircha yozuv yo'q."
    else:
        for e in log[-15:][::-1]:
            text += (
                f"🕒 {_fmt_dt(e.get('created_at',''))}\n"
                f"   👤 {e.get('admin_name','—')} [{e.get('role','')}]\n"
                f"   ▫️ {e.get('action','—')}: {e.get('target','')}\n\n"
            )
    kb = InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="📄 Word (.docx) yuklab olish", callback_data="admin_log_word")],
        [InlineKeyboardButton(text="🔙 Orqaga", callback_data="admin_back")],
    ])
    await _admin_nav(call, text, kb)
    await call.answer()


def _build_log_docx(log: list) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    doc = Document()
    h = doc.add_heading("ProMan Market — Amallar jurnali", level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph(f"Yuklab olingan sana: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    table = doc.add_table(rows=1, cols=5)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, t in enumerate(["#", "Sana / vaqt", "Admin", "Amal", "Tafsilot"]):
        hdr[i].text = t
        for r in hdr[i].paragraphs[0].runs:
            r.font.bold = True

    for idx, e in enumerate(log, 1):
        c = table.add_row().cells
        c[0].text = str(idx)
        c[1].text = _fmt_dt(e.get("created_at", ""))
        c[2].text = f"{e.get('admin_name','—')} [{e.get('role','')}]"
        c[3].text = e.get("action", "—")
        c[4].text = e.get("target", "")

    doc.add_paragraph()
    doc.add_paragraph(f"Jami amallar: {len(log)} ta")
    import io as _io
    buf = _io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


@router.callback_query(F.data == "admin_log_word")
async def admin_log_word(call: CallbackQuery):
    if not is_owner(call.from_user.id): return
    log = get_audit()
    if not log:
        await call.answer("Jurnal bo'sh.", show_alert=True); return
    await call.answer("⏳ Tayyorlanmoqda...")
    data = _build_log_docx(log)
    fname = f"jurnal_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"
    file = BufferedInputFile(data, filename=fname)
    await call.message.answer_document(file, caption=f"📜 Amallar jurnali — {len(log)} ta yozuv")


def _build_sellers_docx(sellers: list) -> bytes:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    doc = Document()
    h = doc.add_heading("ProMan Market — Sellerlar ro'yxati", level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph(f"Yuklab olingan sana: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    table = doc.add_table(rows=1, cols=6)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, t in enumerate(["#", "Ism-familiya", "Telefon", "Shahar", "Do'kon", "ID"]):
        hdr[i].text = t
        for r in hdr[i].paragraphs[0].runs:
            r.font.bold = True

    for idx, s in enumerate(sellers, 1):
        c = table.add_row().cells
        c[0].text = str(idx)
        c[1].text = str(s.get("full_name", "—"))
        c[2].text = str(s.get("phone", "—"))
        c[3].text = str(s.get("city", "—"))
        c[4].text = str(s.get("shop_name", "—"))
        c[5].text = str(s.get("user_id", "—"))

    doc.add_paragraph()
    doc.add_paragraph(f"Jami sellerlar: {len(sellers)} ta")
    import io as _io
    buf = _io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


@router.callback_query(F.data == "admin_sellers_word")
async def admin_sellers_word(call: CallbackQuery):
    if not is_owner(call.from_user.id): return
    sellers = list(get_sellers().values())
    if not sellers:
        await call.answer("Sellerlar yo'q.", show_alert=True); return
    await call.answer("⏳ Tayyorlanmoqda...")
    data = _build_sellers_docx(sellers)
    fname = f"sellerlar_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"
    file = BufferedInputFile(data, filename=fname)
    await call.message.answer_document(file, caption=f"📄 Sellerlar ro'yxati — {len(sellers)} ta")


# ═══════════════════════════════════════════════════════════════════════════
#  OWNER: SHAHARLAR (tuman/shahar) BOSHQARUVI
# ═══════════════════════════════════════════════════════════════════════════
class AdminCityState(StatesGroup):
    name = State()


def _cities_kb():
    rows = []
    for c in get_cities():
        rows.append([
            InlineKeyboardButton(text=f"🏙 {c}", callback_data="noop"),
            InlineKeyboardButton(text="🗑", callback_data=f"delcity_{c}"),
        ])
    rows.append([InlineKeyboardButton(text="➕ Shahar qo'shish", callback_data="addcity")])
    rows.append([InlineKeyboardButton(text="🔙 Orqaga", callback_data="admin_back")])
    return InlineKeyboardMarkup(inline_keyboard=rows)


@router.callback_query(F.data == "admin_cities")
async def admin_cities(call: CallbackQuery):
    if not is_owner(call.from_user.id): return
    await _ack(call)
    cities = get_cities()
    text = (f"🏙 <b>Shaharlar / Tumanlar</b> — {len(cities)} ta\n\n"
            "Sellerlar va xaridorlar shu ro'yxatdan tanlaydi.\n"
            "Har bir shaharga alohida admin tayinlashingiz mumkin.")
    await _admin_nav(call, text, _cities_kb())


@router.callback_query(F.data == "noop")
async def _noop(call: CallbackQuery):
    await call.answer()


@router.callback_query(F.data == "addcity")
async def add_city_start(call: CallbackQuery, state: FSMContext):
    if not is_owner(call.from_user.id): return
    await state.set_state(AdminCityState.name)
    await call.message.answer("🏙 Yangi shahar/tuman nomini yozing:")
    await call.answer()


@router.message(AdminCityState.name)
async def add_city_save(message: Message, state: FSMContext):
    if not is_owner(message.from_user.id):
        await state.clear(); return
    name = (message.text or "").strip()
    await state.clear()
    if not name:
        await message.answer("❌ Nomi bo'sh bo'lmasin."); return
    if add_city(name):
        _log(message.from_user, "Shahar qo'shildi", name)
        await message.answer(f"✅ «{name}» qo'shildi.", reply_markup=admin_menu_kb(message.from_user.id))
    else:
        await message.answer("Bu shahar allaqachon bor.", reply_markup=admin_menu_kb(message.from_user.id))


@router.callback_query(F.data.startswith("delcity_"))
async def del_city_cb(call: CallbackQuery):
    if not is_owner(call.from_user.id): return
    name = call.data.split("_", 1)[1]
    if remove_city(name):
        _log(call.from_user, "Shahar o'chirildi", name)
        await call.answer(f"{name} o'chirildi.")
    else:
        await call.answer("Topilmadi.")
    await admin_cities(call)


# ═══════════════════════════════════════════════════════════════════════════
#  OWNER: ZAXIRA (BACKUP) VA TIKLASH (RESTORE)
# ═══════════════════════════════════════════════════════════════════════════
from app.storage import DATA_DIR


class AdminRestore(StatesGroup):
    waiting_file = State()


def _make_backup_zip() -> tuple[bytes, int]:
    import glob
    files = glob.glob(os.path.join(DATA_DIR, "*.json"))
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as z:
        for f in files:
            z.write(f, arcname=os.path.basename(f))
    return buf.getvalue(), len(files)


async def _send_backup(message: Message):
    data, n = _make_backup_zip()
    if n == 0:
        await message.answer("Hozircha saqlangan ma'lumot yo'q.")
        return
    fname = f"backup_{datetime.now().strftime('%Y%m%d_%H%M')}.zip"
    await message.answer_document(
        BufferedInputFile(data, filename=fname),
        caption=(
            f"💾 <b>Zaxira nusxa</b> — {n} ta fayl\n"
            f"📅 {datetime.now().strftime('%d.%m.%Y %H:%M')}\n\n"
            "Bu faylni xavfsiz joyda saqlang. Yangi serverga ko'chsangiz, "
            "shu faylni botga /restore bilan yuborib hamma narsani tiklaysiz."
        ),
        parse_mode="HTML"
    )


@router.message(Command("backup"))
async def backup_cmd(message: Message):
    if not is_owner(message.from_user.id): return
    await _send_backup(message)


@router.callback_query(F.data == "admin_backup")
async def backup_btn(call: CallbackQuery):
    if not is_owner(call.from_user.id): return
    await call.answer("⏳ Zaxira tayyorlanmoqda...")
    await _send_backup(call.message)


@router.message(Command("restore"))
async def restore_cmd(message: Message, state: FSMContext):
    if not is_owner(message.from_user.id): return
    await state.set_state(AdminRestore.waiting_file)
    await message.answer(
        "♻️ <b>Ma'lumotlarni tiklash</b>\n\n"
        "Avval olingan <b>backup .zip</b> faylini shu yerga (hujjat sifatida) yuboring.\n"
        "⚠️ Diqqat: hozirgi ma'lumotlar yangisi bilan almashtiriladi.\n\n"
        "Bekor qilish: /cancel",
        parse_mode="HTML"
    )


@router.message(AdminRestore.waiting_file, F.text == "/cancel")
async def restore_cancel(message: Message, state: FSMContext):
    await state.clear()
    await message.answer("Bekor qilindi.", reply_markup=admin_menu_kb(message.from_user.id))


@router.message(AdminRestore.waiting_file, F.document)
async def restore_file(message: Message, state: FSMContext):
    if not is_owner(message.from_user.id):
        await state.clear(); return
    doc = message.document
    if not (doc.file_name or "").endswith(".zip"):
        await message.answer("❌ Iltimos, backup .zip faylini yuboring.")
        return
    await state.clear()
    try:
        from app.bot.bot import bot
        bio = io.BytesIO()
        await bot.download(doc, destination=bio)
        bio.seek(0)
        os.makedirs(DATA_DIR, exist_ok=True)
        count = 0
        with zipfile.ZipFile(bio) as z:
            for n in z.namelist():
                if not n.endswith(".json"):
                    continue
                base = os.path.basename(n)        # xavfsizlik: papka yo'llarini tashlaymiz
                if not base:
                    continue
                with z.open(n) as src, open(os.path.join(DATA_DIR, base), "wb") as dst:
                    dst.write(src.read())
                count += 1
        if count == 0:
            await message.answer("❌ Zip ichida .json fayl topilmadi. To'g'ri backup faylini yuboring.")
            return
        _log(message.from_user, "Ma'lumotlar tiklandi", f"{count} ta fayl")
        await message.answer(
            f"✅ <b>{count} ta fayl tiklandi!</b>\nMa'lumotlar muvaffaqiyatli yangilandi.",
            parse_mode="HTML", reply_markup=admin_menu_kb(message.from_user.id)
        )
    except Exception as e:
        await message.answer(f"❌ Faylni o'qishda xato: {e}")


@router.message(AdminRestore.waiting_file)
async def restore_wrong(message: Message):
    await message.answer("❌ Backup .zip faylini hujjat sifatida yuboring (yoki /cancel).")
